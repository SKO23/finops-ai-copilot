import os
import re
import html
import json
import time
import random
import csv
import tempfile
from datetime import datetime
from pathlib import Path

import numpy as np
import gradio as gr
from pypdf import PdfReader
from dotenv import load_dotenv
from google import genai
from google.genai import types
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# ---------------------------------------------------------
# Configuration
# ---------------------------------------------------------

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash")
GEMINI_FALLBACK_MODEL = os.getenv(
    "GEMINI_FALLBACK_MODEL",
    "gemini-3.5-flash-lite",
)
GEMINI_MAX_ATTEMPTS = 3
GEMINI_BUSY_MESSAGE = (
    "Gemini is temporarily busy. Your documents were processed successfully "
    "and remain ready. Please try the same action again in a moment."
)

MIN_RELEVANCE_SCORE = 0.30
TOP_K = 6
MAX_SUBQUESTIONS = 5

IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}
MIN_NATIVE_TEXT_CHARS = 80

client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

embedder = SentenceTransformer(
    "sentence-transformers/all-MiniLM-L6-v2"
)

EXPORT_DIR = Path(tempfile.gettempdir()) / "finops_ai_copilot_exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

COMPLETE_REVIEW_INSTRUCTIONS = {
    "risks": (
        "Extract the travel or business expense items from the uploaded documents. "
        "For each item identify expense type, date, amount, currency, merchant or provider, "
        "claimant where stated, and the supporting source/page. Do not invent missing values."
    ),
    "actions": (
        "Check each identifiable expense against the uploaded finance or travel policy. "
        "State the relevant policy rule or limit, whether the expense is Compliant, Exception, "
        "or Needs Review, and the supporting evidence. If no policy limit is stated, say Not stated."
    ),
    "dates": (
        "Identify exceptions, policy breaches, unusual items, duplicate-looking claims, unsupported "
        "amounts, or other finance-review risks that are supported by the uploaded documents. "
        "Assign High, Medium or Low severity and recommend the next finance action."
    ),
    "comparison": (
        "Identify missing evidence needed to complete the finance review, such as a receipt, invoice, "
        "approval, business purpose, attendee details, travel justification, or policy evidence. "
        "Only mark evidence as missing when the uploaded documents do not provide it."
    ),
}


# ---------------------------------------------------------
# Reusable UI content
# ---------------------------------------------------------

EMPTY_SUMMARY_HTML = """
<div class="summary-box empty-panel">
    <div class="empty-icon">📝</div>
    <strong>No summary yet</strong>
    <span>Upload and process expense claims, receipts and policy documents to create a finance review summary.</span>
</div>
"""

EMPTY_ANSWER_HTML = """
<div class="answer-box empty-panel">
    <div class="empty-icon">💬</div>
    <strong>No answer yet</strong>
    <span>Ask a finance or policy question after your documents have been processed.</span>
</div>
"""

EMPTY_RECOMMENDATION_HTML = """
<div class="answer-box empty-panel">
    <div class="empty-icon">🤖</div>
    <strong>No AI recommendation yet</strong>
    <span>Generate the advisory recommendation after processing the finance documents.</span>
</div>
"""

NOT_FOUND_MESSAGE = (
    "I could not find this information in the uploaded documents."
)

STRUCTURED_REVIEW_CONFIG = {
    "risks": {
        "title": "Expense Summary",
        "schema": {
            "summary": "Short overall expense summary",
            "items": [
                {
                    "expense_type": "Expense category or type",
                    "date": "Expense date or Not stated",
                    "amount": "Expense amount exactly as shown",
                    "currency": "Currency or Not stated",
                    "merchant": "Merchant, supplier or provider, or Not stated",
                    "claimant": "Employee or claimant, or Not stated",
                    "source": "Exact uploaded file name",
                    "page": "Page number as a string",
                }
            ],
        },
        "columns": [
            ("Expense Type", "expense_type"),
            ("Date", "date"),
            ("Amount", "amount"),
            ("Currency", "currency"),
            ("Merchant / Provider", "merchant"),
            ("Claimant", "claimant"),
            ("Source", "source"),
            ("Page", "page"),
        ],
    },
    "actions": {
        "title": "Policy Compliance",
        "schema": {
            "summary": "Short overall policy-compliance summary",
            "items": [
                {
                    "expense_type": "Expense item or category being checked",
                    "amount": "Claimed amount exactly as shown",
                    "policy_limit": "Relevant policy rule or monetary limit, or Not stated",
                    "status": "Compliant, Exception or Needs Review",
                    "evidence": "Short explanation grounded in the uploaded documents",
                    "source": "Exact supporting uploaded file name",
                    "page": "Supporting page number as a string",
                }
            ],
        },
        "columns": [
            ("Expense Type", "expense_type"),
            ("Amount", "amount"),
            ("Policy Limit / Rule", "policy_limit"),
            ("Status", "status"),
            ("Evidence", "evidence"),
            ("Source", "source"),
            ("Page", "page"),
        ],
    },
    "dates": {
        "title": "Exceptions & Risks",
        "schema": {
            "summary": "Short overall exceptions and risk summary",
            "items": [
                {
                    "severity": "High, Medium or Low",
                    "finding": "Exception, policy issue or finance-review risk",
                    "reason": "Why it is an exception or risk based on the documents",
                    "recommended_action": "Recommended finance follow-up",
                    "source": "Exact supporting uploaded file name",
                    "page": "Supporting page number as a string",
                }
            ],
        },
        "columns": [
            ("Severity", "severity"),
            ("Exception / Risk", "finding"),
            ("Reason", "reason"),
            ("Recommended Action", "recommended_action"),
            ("Source", "source"),
            ("Page", "page"),
        ],
    },
    "comparison": {
        "title": "Missing Evidence",
        "schema": {
            "summary": "Short overall missing-evidence summary",
            "items": [
                {
                    "expense_or_claim": "Expense or claim affected",
                    "missing_evidence": "Receipt, invoice, approval, business purpose or other missing evidence",
                    "impact": "Why the missing evidence matters for finance review",
                    "recommended_action": "What the reviewer should request or check next",
                    "source": "Exact related uploaded file name, or Not stated",
                    "page": "Related page number as a string, or Not stated",
                }
            ],
        },
        "columns": [
            ("Expense / Claim", "expense_or_claim"),
            ("Missing Evidence", "missing_evidence"),
            ("Review Impact", "impact"),
            ("Recommended Action", "recommended_action"),
            ("Source", "source"),
            ("Page", "page"),
        ],
    },
}


# ---------------------------------------------------------
# Text extraction and chunking
# ---------------------------------------------------------

def clean_text(text):
    """Remove repeated whitespace from extracted PDF text."""
    return re.sub(r"\s+", " ", text or "").strip()


def extract_pdf_pages(pdf_path):
    """Extract selectable PDF text page by page and preserve source/page metadata."""
    reader = PdfReader(pdf_path)
    source_name = Path(pdf_path).name
    extracted_pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = clean_text(page.extract_text() or "")

        if page_text:
            extracted_pages.append(
                {
                    "source": source_name,
                    "page": page_number,
                    "text": page_text,
                }
            )

    return extracted_pages, len(reader.pages)


def call_gemini_with_binary(prompt, file_path, mime_type):
    """Send a PDF/image to Gemini for vision-based text extraction."""
    if not client:
        return None, (
            "GEMINI_API_KEY is missing. Add it to Hugging Face Space Secrets "
            "before using scanned-document extraction."
        )

    try:
        file_bytes = Path(file_path).read_bytes()
        file_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
    except Exception:
        return None, "The uploaded file could not be prepared for Gemini extraction."

    models_to_try = [GEMINI_MODEL]
    if GEMINI_FALLBACK_MODEL and GEMINI_FALLBACK_MODEL != GEMINI_MODEL:
        models_to_try.append(GEMINI_FALLBACK_MODEL)

    for model_index, model_name in enumerate(models_to_try):
        for attempt in range(GEMINI_MAX_ATTEMPTS):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=[file_part, prompt],
                )

                if response.text:
                    return response.text.strip(), None

            except Exception as error:
                if not is_retryable_gemini_error(error):
                    return None, (
                        "Gemini could not read this scanned/image document. "
                        "Please check the file and model settings."
                    )

            if attempt < GEMINI_MAX_ATTEMPTS - 1:
                delay_seconds = (1.5 * (2 ** attempt)) + random.uniform(0, 0.5)
                time.sleep(delay_seconds)

        if model_index < len(models_to_try) - 1:
            continue

    return None, GEMINI_BUSY_MESSAGE


def parse_page_marked_transcription(transcription, source_name):
    """Convert Gemini page-marked transcription into normal page records."""
    transcription = (transcription or "").strip()
    if not transcription:
        return []

    transcription = re.sub(r"^```(?:text|markdown)?\s*", "", transcription, flags=re.IGNORECASE)
    transcription = re.sub(r"\s*```$", "", transcription).strip()

    page_pattern = re.compile(
        r"===\s*PAGE\s*(\d+)\s*===\s*(.*?)(?====\s*PAGE\s*\d+\s*===|\Z)",
        flags=re.IGNORECASE | re.DOTALL,
    )
    matches = page_pattern.findall(transcription)
    pages = []

    if matches:
        for page_number, page_text in matches:
            cleaned = clean_text(page_text)
            if cleaned and cleaned.upper() != "[UNREADABLE]":
                pages.append(
                    {
                        "source": source_name,
                        "page": int(page_number),
                        "text": cleaned,
                    }
                )
        return pages

    cleaned = clean_text(transcription)
    if cleaned:
        return [{"source": source_name, "page": 1, "text": cleaned}]

    return []


def extract_with_gemini_vision(file_path, mime_type):
    """Use Gemini native document/image vision as an OCR-style fallback."""
    source_name = Path(file_path).name
    prompt = """
You are a document transcription engine for a finance-review application.

Transcribe ALL visible text from this uploaded document faithfully.
Do not summarize, interpret, correct, infer or add information.
Preserve amounts, currencies, dates, merchant names, invoice/receipt numbers,
employee names, policy limits, approvals, table values and handwritten text when readable.

Separate every PDF page using exactly this marker:
=== PAGE 1 ===
<transcribed text>
=== PAGE 2 ===
<transcribed text>

For a single image, use exactly:
=== PAGE 1 ===
<transcribed text>

If a page is genuinely unreadable, write [UNREADABLE] under that page marker.
Return only the page-marked transcription, with no commentary and no code fence.
"""
    transcription, error = call_gemini_with_binary(prompt, file_path, mime_type)
    if error:
        return [], error

    pages = parse_page_marked_transcription(transcription, source_name)
    if not pages:
        return [], "Gemini did not return readable text from this document."

    return pages, None


def split_page_into_chunks(page_data, chunk_size=350, overlap=60):
    """Split one page into smaller searchable chunks."""
    words = page_data["text"].split()
    chunks = []

    if not words:
        return chunks

    start = 0
    chunk_number = 1

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_text = " ".join(words[start:end])

        chunks.append(
            {
                "text": chunk_text,
                "source": page_data["source"],
                "page": page_data["page"],
                "chunk_number": chunk_number,
            }
        )

        if end >= len(words):
            break

        start = max(end - overlap, start + 1)
        chunk_number += 1

    return chunks


def build_chunks(all_pages):
    """Create searchable chunks from all readable pages."""
    chunks = []

    for page_data in all_pages:
        chunks.extend(split_page_into_chunks(page_data))

    return chunks


# ---------------------------------------------------------
# Embeddings and retrieval
# ---------------------------------------------------------

def embed_chunks(chunks):
    """Create an embedding for every document chunk."""
    if not chunks:
        return None

    chunk_texts = [chunk["text"] for chunk in chunks]

    embeddings = embedder.encode(
        chunk_texts,
        convert_to_numpy=True,
        normalize_embeddings=True,
        show_progress_bar=False,
    )

    return embeddings.astype("float32")


SEARCH_STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "been", "by",
    "can", "could", "did", "do", "does", "for", "from", "has",
    "have", "how", "i", "in", "is", "it", "me", "of", "on",
    "or", "please", "shown", "tell", "that", "the", "their",
    "there", "this", "to", "was", "were", "what", "when",
    "where", "which", "who", "why", "will", "with", "would",
}

ACRONYM_LINK_WORDS = {
    "a", "an", "and", "for", "in", "of", "on", "the", "to",
}


def normalise_for_search(text):
    """Normalise text for exact and phrase matching."""
    text = text or ""

    # Convert spaced acronyms such as "S V R" into "SVR".
    text = re.sub(
        r"(?<![A-Za-z0-9])(?:[A-Za-z]\s+){1,7}[A-Za-z](?![A-Za-z0-9])",
        lambda match: re.sub(r"\s+", "", match.group(0)),
        text,
    )

    text = text.lower()
    text = re.sub(r"[^a-z0-9£$%]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def extract_search_terms(question):
    """Extract useful query terms while removing common question words."""
    terms = []

    for term in normalise_for_search(question).split():
        if term in SEARCH_STOP_WORDS:
            continue

        if len(term) < 2 and not term.isdigit():
            continue

        if term not in terms:
            terms.append(term)

    return terms


def extract_acronym_terms(question):
    """Detect genuine uppercase acronyms in the user's wording."""
    question = question or ""
    acronyms = []

    for token in re.findall(r"\b[A-Z][A-Z0-9]{1,7}\b", question):
        token_lower = token.lower()

        if token_lower not in acronyms:
            acronyms.append(token_lower)

    for match in re.findall(
        r"(?<![A-Za-z0-9])(?:[A-Z]\s+){1,7}[A-Z](?![A-Za-z0-9])",
        question,
    ):
        collapsed = re.sub(r"\s+", "", match).lower()

        if collapsed not in acronyms:
            acronyms.append(collapsed)

    return acronyms


def acronym_initials(words):
    """Create initials while ignoring small linking words."""
    significant_words = [
        word
        for word in words
        if word.lower() not in ACRONYM_LINK_WORDS
    ]

    return "".join(word[0] for word in significant_words).lower()


def discover_acronym_expansions(question, chunks):
    """
    Discover possible acronym expansions directly from uploaded documents.

    Example: SVR can match the document phrase "Standard Variable Rate"
    even when the letters "SVR" never appear in the PDF.
    """
    acronyms = extract_acronym_terms(question)
    discovered = {}

    if not acronyms:
        return discovered

    for acronym in acronyms:
        candidates = {}

        for chunk in chunks:
            text = chunk.get("text", "")
            source = chunk.get("source", "")
            page = chunk.get("page", 0)

            # Strong pattern: Full Phrase (ACR) or ACR (Full Phrase).
            flexible_acronym = make_spaced_acronym_pattern(acronym)
            parenthetical_patterns = [
                rf"([A-Za-z][A-Za-z'\-]*(?:\s+[A-Za-z][A-Za-z'\-]*){{1,8}})\s*\(\s*{flexible_acronym}\s*\)",
                rf"\b{flexible_acronym}\b\s*\(\s*([A-Za-z][A-Za-z'\-]*(?:\s+[A-Za-z][A-Za-z'\-]*){{1,8}})\s*\)",
            ]

            for pattern in parenthetical_patterns:
                for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                    phrase = re.sub(r"\s+", " ", match.group(1)).strip()
                    words = re.findall(r"[A-Za-z][A-Za-z'\-]*", phrase)

                    if acronym_initials(words) == acronym:
                        key = normalise_for_search(phrase)
                        candidates.setdefault(
                            key,
                            {
                                "phrase": phrase,
                                "score": 0.0,
                                "sources": set(),
                            },
                        )
                        candidates[key]["score"] += 10.0
                        candidates[key]["sources"].add((source, page))

            # General pattern: title-case phrase whose initials equal the acronym.
            words_with_case = re.findall(r"[A-Za-z][A-Za-z'\-]*", text)
            target_length = len(acronym)

            for start in range(len(words_with_case)):
                for window_size in range(target_length, target_length + 3):
                    window = words_with_case[start:start + window_size]

                    if len(window) < target_length:
                        continue

                    significant = [
                        word
                        for word in window
                        if word.lower() not in ACRONYM_LINK_WORDS
                    ]

                    if len(significant) != target_length:
                        continue

                    if acronym_initials(window) != acronym:
                        continue

                    # Requiring title case avoids random matches in normal prose.
                    if not all(word[0].isupper() for word in significant):
                        continue

                    phrase = " ".join(window)
                    key = normalise_for_search(phrase)

                    candidates.setdefault(
                        key,
                        {
                            "phrase": phrase,
                            "score": 0.0,
                            "sources": set(),
                        },
                    )
                    candidates[key]["score"] += 4.0
                    candidates[key]["sources"].add((source, page))

        ranked = sorted(
            candidates.values(),
            key=lambda item: (
                item["score"],
                len(item["sources"]),
                -len(item["phrase"]),
            ),
            reverse=True,
        )

        if ranked:
            best = ranked[0]
            discovered[acronym] = {
                "phrase": best["phrase"],
                "sources": sorted(best["sources"]),
            }

    return discovered


def calculate_keyword_match(question, chunk_text):
    """Calculate exact-term coverage for a document chunk."""
    question_terms = extract_search_terms(question)
    chunk_terms = set(normalise_for_search(chunk_text).split())

    matched_terms = [
        term for term in question_terms
        if term in chunk_terms
    ]

    term_coverage = (
        len(matched_terms) / len(question_terms)
        if question_terms
        else 0.0
    )

    meaningful_phrase = " ".join(question_terms)
    phrase_match = (
        len(question_terms) >= 2
        and meaningful_phrase in normalise_for_search(chunk_text)
    )

    keyword_score = min(
        term_coverage + (0.20 if phrase_match else 0.0),
        1.0,
    )

    return {
        "keyword_score": keyword_score,
        "term_coverage": term_coverage,
        "matched_terms": matched_terms,
    }


def retrieve_relevant_chunks(
    question,
    chunks,
    embeddings,
    top_k=TOP_K,
    min_score=MIN_RELEVANCE_SCORE,
):
    """
    Retrieve evidence using semantic search, exact keywords and
    document-derived acronym expansion.
    """
    if not chunks or embeddings is None:
        return [], {"acronym_expansions": {}}

    acronym_expansions = discover_acronym_expansions(question, chunks)
    acronym_terms = extract_acronym_terms(question)

    search_variants = [question]

    for acronym, details in acronym_expansions.items():
        expansion_query = f"{question} {details['phrase']}"

        if expansion_query not in search_variants:
            search_variants.append(expansion_query)

    query_embeddings = embedder.encode(
        search_variants,
        convert_to_numpy=True,
        normalize_embeddings=True,
        show_progress_bar=False,
    ).astype("float32")

    # Use the best semantic score across the original and expanded queries.
    semantic_matrix = np.dot(embeddings, query_embeddings.T)
    best_semantic_scores = np.max(semantic_matrix, axis=1)

    ranked_results = []

    for index, chunk in enumerate(chunks):
        semantic_score = float(best_semantic_scores[index])
        keyword_details = calculate_keyword_match(question, chunk["text"])
        normalised_chunk = normalise_for_search(chunk["text"])
        chunk_terms = set(normalised_chunk.split())

        exact_acronym_matches = [
            acronym
            for acronym in acronym_terms
            if acronym in chunk_terms
        ]

        expansion_matches = []

        for acronym, details in acronym_expansions.items():
            expansion_text = normalise_for_search(details["phrase"])

            if expansion_text and expansion_text in normalised_chunk:
                expansion_matches.append(
                    {
                        "acronym": acronym,
                        "phrase": details["phrase"],
                    }
                )

        exact_acronym_match = bool(exact_acronym_matches)
        expansion_match = bool(expansion_matches)

        hybrid_score = (
            semantic_score
            + (0.30 * keyword_details["keyword_score"])
            + (1.00 if exact_acronym_match else 0.0)
            + (0.90 if expansion_match else 0.0)
        )

        is_relevant = (
            exact_acronym_match
            or expansion_match
            or semantic_score >= min_score
            or (
                keyword_details["term_coverage"] >= 0.60
                and semantic_score >= 0.12
            )
        )

        if not is_relevant:
            continue

        ranked_results.append(
            {
                "index": index,
                "hybrid_score": hybrid_score,
                "semantic_score": semantic_score,
                "exact_acronym_match": exact_acronym_match,
                "exact_acronym_matches": exact_acronym_matches,
                "expansion_match": expansion_match,
                "expansion_matches": expansion_matches,
                **keyword_details,
            }
        )

    ranked_results.sort(
        key=lambda item: (
            1 if item["expansion_match"] else 0,
            1 if item["exact_acronym_match"] else 0,
            item["hybrid_score"],
        ),
        reverse=True,
    )

    selected_chunks = []
    seen = set()

    for result in ranked_results:
        chunk = chunks[result["index"]]
        unique_key = (
            chunk["source"],
            chunk["page"],
            chunk["text"][:180],
        )

        if unique_key in seen:
            continue

        seen.add(unique_key)

        selected_chunks.append(
            {
                **chunk,
                "score": float(result["hybrid_score"]),
                "semantic_score": float(result["semantic_score"]),
                "keyword_score": float(result["keyword_score"]),
                "matched_terms": result["matched_terms"],
                "exact_acronym_match": result["exact_acronym_match"],
                "exact_acronym_matches": result["exact_acronym_matches"],
                "expansion_match": result["expansion_match"],
                "expansion_matches": result["expansion_matches"],
            }
        )

        if len(selected_chunks) >= min(top_k, len(chunks)):
            break

    return selected_chunks, {
        "acronym_expansions": acronym_expansions,
        "search_variants": search_variants,
    }

# ---------------------------------------------------------
# Gemini helpers
# ---------------------------------------------------------

def is_retryable_gemini_error(error):
    """Return True for temporary Gemini capacity or rate-limit errors."""
    error_text = str(error).lower()

    retryable_signals = (
        "503",
        "unavailable",
        "high demand",
        "overloaded",
        "resource_exhausted",
        "429",
        "temporarily",
        "deadline exceeded",
        "timeout",
    )

    return any(signal in error_text for signal in retryable_signals)


def call_gemini(prompt):
    """
    Send a prompt to Gemini with exponential-backoff retries.

    The primary model is tried first. If temporary capacity errors continue,
    the app automatically tries the configured Flash-Lite fallback model.
    """
    if not client:
        return (
            "Error: GEMINI_API_KEY is missing. Add it to your local .env "
            "file or to Hugging Face Space Secrets."
        )

    models_to_try = [GEMINI_MODEL]

    if (
        GEMINI_FALLBACK_MODEL
        and GEMINI_FALLBACK_MODEL != GEMINI_MODEL
    ):
        models_to_try.append(GEMINI_FALLBACK_MODEL)

    last_error = None

    for model_index, model_name in enumerate(models_to_try):
        for attempt in range(GEMINI_MAX_ATTEMPTS):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                )

                if response.text:
                    return response.text.strip()

                last_error = RuntimeError(
                    f"{model_name} returned an empty response."
                )

            except Exception as error:
                last_error = error

                if not is_retryable_gemini_error(error):
                    return (
                        "Gemini could not complete this request. "
                        "Please check the model setting and try again."
                    )

            is_last_attempt = attempt == GEMINI_MAX_ATTEMPTS - 1

            if not is_last_attempt:
                # 1.5s, 3s, then 6s if more attempts are added, plus jitter.
                delay_seconds = (1.5 * (2 ** attempt)) + random.uniform(0, 0.5)
                time.sleep(delay_seconds)

        # Move immediately to the fallback model after the primary retries.
        if model_index < len(models_to_try) - 1:
            continue

    # Do not expose the full provider error to end users.
    return GEMINI_BUSY_MESSAGE

# ---------------------------------------------------------
# Multi-question detection
# ---------------------------------------------------------

def normalise_question(text):
    """Clean a question and ensure it ends with a question mark."""
    cleaned = re.sub(r"\s+", " ", text or "").strip(" \t\r\n,;.-")

    if cleaned and not cleaned.endswith("?"):
        cleaned += "?"

    return cleaned


def split_list_style_question(question):
    """
    Split list-style requests such as:
    "What is the interest rate, definition of SVR and loan offered?"
    """
    cleaned = re.sub(r"\s+", " ", question or "").strip().rstrip("?")

    match = re.match(
        r"^(what|which|who|where|when|why|how)\s+"
        r"(is|are|was|were|does|do|did|can|could|should|would|will|has|have)"
        r"\s+(?:the\s+)?(.+)$",
        cleaned,
        flags=re.IGNORECASE,
    )

    if not match:
        return []

    question_word, helper_verb, remainder = match.groups()

    if "," not in remainder or not re.search(
        r"\s+(?:and|also|plus|as well as)\s+",
        remainder,
        flags=re.IGNORECASE,
    ):
        return []

    items = re.split(
        r"\s*,\s*|\s+(?:and|also|plus|as well as)\s+",
        remainder,
        flags=re.IGNORECASE,
    )

    items = [item.strip(" ,;.-") for item in items if item.strip(" ,;.-")]

    if len(items) < 2:
        return []

    questions = []

    for item in items[:MAX_SUBQUESTIONS]:
        item_lower = item.lower()

        if item_lower.startswith(
            ("what ", "which ", "who ", "where ", "when ", "why ", "how ")
        ):
            questions.append(normalise_question(item))
        else:
            article = "" if item_lower.startswith("the ") else "the "
            questions.append(
                normalise_question(
                    f"{question_word} {helper_verb} {article}{item}"
                )
            )

    return questions


def split_into_subquestions(question):
    """Split multi-part requests into independent similarity searches."""
    original = re.sub(r"\s+", " ", question or "").strip()

    if not original:
        return []

    list_questions = split_list_style_question(original)

    if list_questions:
        return list_questions

    broad_parts = re.split(r"\s*;\s*|\s*\n+\s*|(?<=\?)\s+", original)
    parts = []

    connector_pattern = re.compile(
        r"\s+(?:and|also|plus|as well as)\s+"
        r"(?=(?:what|which|who|where|when|why|how|"
        r"is|are|was|were|do|does|did|can|could|"
        r"should|would|will|has|have)\b)",
        flags=re.IGNORECASE,
    )

    for broad_part in broad_parts:
        for part in connector_pattern.split(broad_part):
            cleaned_part = normalise_question(part)

            if cleaned_part:
                parts.append(cleaned_part)

    unique_parts = []
    seen = set()

    for part in parts:
        key = part.lower()

        if key not in seen:
            seen.add(key)
            unique_parts.append(part)

    return (unique_parts or [normalise_question(original)])[:MAX_SUBQUESTIONS]

# ---------------------------------------------------------
# Summary generation
# ---------------------------------------------------------

def format_ai_text(text):
    """Convert common Gemini Markdown into safe HTML."""
    safe_text = html.escape(text or "")
    lines = safe_text.splitlines()
    rendered = []
    in_ul = False
    in_ol = False

    def close_lists():
        nonlocal in_ul, in_ol

        if in_ul:
            rendered.append("</ul>")
            in_ul = False

        if in_ol:
            rendered.append("</ol>")
            in_ol = False

    def inline_format(line):
        line = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", line)
        line = re.sub(
            r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)",
            r"<em>\1</em>",
            line,
        )
        return line

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            close_lists()
            rendered.append('<div class="text-spacer"></div>')
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)

        if heading:
            close_lists()
            level = min(len(heading.group(1)) + 2, 5)
            rendered.append(
                f"<h{level}>{inline_format(heading.group(2))}</h{level}>"
            )
            continue

        bullet = re.match(r"^[*\-•]\s+(.+)$", line)

        if bullet:
            if in_ol:
                rendered.append("</ol>")
                in_ol = False

            if not in_ul:
                rendered.append("<ul>")
                in_ul = True

            rendered.append(f"<li>{inline_format(bullet.group(1))}</li>")
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)

        if numbered:
            if in_ul:
                rendered.append("</ul>")
                in_ul = False

            if not in_ol:
                rendered.append("<ol>")
                in_ol = True

            rendered.append(f"<li>{inline_format(numbered.group(1))}</li>")
            continue

        close_lists()
        rendered.append(f"<p>{inline_format(line)}</p>")

    close_lists()
    return "".join(rendered)


def make_summary_html(summary_text):
    if summary_text == GEMINI_BUSY_MESSAGE:
        return f"""
        <div class="summary-box service-warning-panel">
            <div class="panel-heading">
                <span>📝 Finance case summary</span>
            </div>
            <div class="service-warning">
                <strong>⏳ Gemini is temporarily busy</strong>
                <span>{html.escape(GEMINI_BUSY_MESSAGE)}</span>
            </div>
        </div>
        """

    formatted_summary = format_ai_text(summary_text)

    return f"""
    <div class="summary-box">
        <div class="panel-heading">
            <span>📝 Finance case summary</span>
        </div>
        <div class="panel-content formatted-ai-text">
            {formatted_summary}
        </div>
    </div>
    """

def create_summary_context(all_pages, max_characters=18000):
    """Build a balanced summary sample across all uploaded PDFs."""
    if not all_pages:
        return ""

    pages_by_source = {}

    for page in all_pages:
        pages_by_source.setdefault(page["source"], []).append(page)

    document_count = len(pages_by_source)
    per_document_limit = max(
        2500,
        max_characters // max(document_count, 1),
    )

    sections = []

    for source_name, pages in pages_by_source.items():
        document_text = []
        current_length = 0

        for page in pages:
            labelled_text = (
                f"[Source: {source_name} | Page {page['page']}]\n"
                f"{page['text']}"
            )

            remaining = per_document_limit - current_length

            if remaining <= 0:
                break

            text_to_add = labelled_text[:remaining]
            document_text.append(text_to_add)
            current_length += len(text_to_add)

        sections.append("\n\n".join(document_text))

    return "\n\n--- NEXT DOCUMENT ---\n\n".join(sections)[
        :max_characters
    ]


def process_documents(document_files):
    """Process PDFs/images, use OCR fallback where needed, embed text and create a summary."""
    if not document_files:
        return (
            "<div class='status-box error'>Please upload at least one finance document first.</div>",
            EMPTY_SUMMARY_HTML,
            [],
            None,
            "",
            [],
        )

    if isinstance(document_files, str):
        document_files = [document_files]

    try:
        all_pages = []
        processed_documents = []
        warnings = []
        total_input_pages = 0

        for document_path in document_files:
            source_name = Path(document_path).name
            suffix = Path(document_path).suffix.lower()

            try:
                extracted_pages = []
                total_pages = 0
                extraction_method = ""

                if suffix == ".pdf":
                    native_pages, total_pages = extract_pdf_pages(document_path)
                    native_character_count = sum(len(item.get("text", "")) for item in native_pages)
                    needs_vision_fallback = (
                        not native_pages
                        or len(native_pages) < total_pages
                        or native_character_count < MIN_NATIVE_TEXT_CHARS
                    )

                    if needs_vision_fallback:
                        vision_pages, vision_error = extract_with_gemini_vision(
                            document_path,
                            "application/pdf",
                        )
                        if vision_pages:
                            extracted_pages = vision_pages
                            extraction_method = "Gemini vision / OCR fallback"
                            warnings.append(
                                f"{source_name}: Gemini vision was used because selectable PDF text "
                                "was missing or incomplete."
                            )
                        else:
                            extracted_pages = native_pages
                            extraction_method = "Native PDF text" if native_pages else "Not readable"
                            if vision_error:
                                warnings.append(f"{source_name}: {vision_error}")
                    else:
                        extracted_pages = native_pages
                        extraction_method = "Native PDF text"

                elif suffix in IMAGE_MIME_TYPES:
                    total_pages = 1
                    extracted_pages, vision_error = extract_with_gemini_vision(
                        document_path,
                        IMAGE_MIME_TYPES[suffix],
                    )
                    extraction_method = "Gemini vision / OCR" if extracted_pages else "Not readable"
                    if vision_error:
                        warnings.append(f"{source_name}: {vision_error}")

                else:
                    warnings.append(
                        f"{source_name}: unsupported file type. Use PDF, PNG, JPG, JPEG or WEBP."
                    )
                    continue

                total_input_pages += total_pages

                if extracted_pages:
                    all_pages.extend(extracted_pages)
                    processed_documents.append(
                        {
                            "source": source_name,
                            "readable_pages": len(extracted_pages),
                            "total_pages": total_pages,
                            "extraction_method": extraction_method,
                        }
                    )
                else:
                    warnings.append(
                        f"{source_name}: no readable text could be extracted."
                    )

            except Exception as document_error:
                warnings.append(f"{source_name}: {str(document_error)}")

        if not all_pages:
            warning_items = "".join(
                f"<li>{html.escape(item)}</li>"
                for item in warnings
            )

            return (
                f"""
                <div class='status-box error'>
                    <strong>No readable text was found.</strong><br>
                    The native extractor and Gemini vision/OCR fallback could not read the uploaded documents.
                    <ul>{warning_items}</ul>
                </div>
                """,
                """
                <div class="summary-box empty-panel">
                    <div class="empty-icon">⚠️</div>
                    <strong>No readable document text</strong>
                    <span>Try a clearer scan/photo or a text-based PDF.</span>
                </div>
                """,
                [],
                None,
                "",
                [],
            )

        chunks = build_chunks(all_pages)
        embeddings = embed_chunks(chunks)
        summary_context = create_summary_context(all_pages)

        document_names = ", ".join(
            item["source"] for item in processed_documents
        )

        summary_prompt = f"""
You are FinOps AI Copilot, supporting a finance reviewer with travel and business expense checks.

Use only the uploaded document text below. Do not add facts that are not present.

Prepare:
1. A short finance case summary
2. The main expense or claim details visible in the documents
3. The relevant finance/travel policy information where present
4. Any obvious exceptions, missing evidence or items needing manual review
5. Do not make the final approval decision; the human finance reviewer retains that responsibility

Uploaded documents:
{document_names}

DOCUMENT TEXT:
{summary_context}
"""

        summary = call_gemini(summary_prompt)

        document_rows = "".join(
            f"""
            <div class="document-row">
                <span class="document-name">📄 {html.escape(item['source'])}</span>
                <span class="page-count">
                    {item['readable_pages']} readable / {item['total_pages']} total pages
                    · {html.escape(item.get('extraction_method', 'Text extraction'))}
                </span>
            </div>
            """
            for item in processed_documents
        )

        warning_html = ""
        if warnings:
            warning_items = "".join(
                f"<li>{html.escape(warning)}</li>"
                for warning in warnings
            )
            warning_html = f"""
            <div class="warning-list">
                <strong>Processing notes</strong>
                <ul>{warning_items}</ul>
            </div>
            """

        status = f"""
        <div class="status-box success">
            <div class="status-title">✅ Finance documents ready for review</div>
            <div class="status-metrics">
                <span><strong>{len(processed_documents)}</strong> documents</span>
                <span><strong>{total_input_pages}</strong> pages</span>
                <span><strong>{len(chunks)}</strong> searchable chunks</span>
            </div>
            <div class="document-results">{document_rows}</div>
            {warning_html}
        </div>
        """

        return (
            status,
            make_summary_html(summary),
            chunks,
            embeddings,
            summary,
            processed_documents,
        )

    except Exception as error:
        return (
            f"""
            <div class="status-box error">
                <strong>Error processing documents:</strong><br>
                {html.escape(str(error))}
            </div>
            """,
            """
            <div class="summary-box empty-panel">
                <div class="empty-icon">⚠️</div>
                <strong>Processing error</strong>
                <span>The finance case summary could not be created.</span>
            </div>
            """,
            [],
            None,
            "",
            [],
        )



# ---------------------------------------------------------
# Question answering and sources
# ---------------------------------------------------------

def build_context(relevant_chunks):
    """Format retrieved chunks for Gemini."""
    context_sections = []

    for index, item in enumerate(relevant_chunks, start=1):
        context_sections.append(
            f"""Evidence {index}
Source: {item['source']}
Page: {item['page']}
Similarity score: {item['score']:.3f}
Text: {item['text']}"""
        )

    return "\n\n---\n\n".join(context_sections)


def format_source_list(relevant_chunks):
    """Create a deduplicated source list for one answer."""
    unique_sources = []
    seen = set()

    for item in relevant_chunks:
        source_key = (item["source"], item["page"])

        if source_key not in seen:
            seen.add(source_key)
            unique_sources.append(source_key)

    if not unique_sources:
        return ""

    source_items = "".join(
        f"""
        <li>
            <span class="source-file">{html.escape(source)}</span>
            <span class="source-page">Page {page}</span>
        </li>
        """
        for source, page in unique_sources
    )

    return f"""
    <div class="sources-box">
        <div class="sources-title">🔎 Supporting sources</div>
        <ul>{source_items}</ul>
    </div>
    """



def make_spaced_acronym_pattern(acronym):
    """Return a regex that matches both SVR and S V R."""
    letters = [re.escape(letter) for letter in acronym]
    return r"\s*".join(letters)


def split_into_sentences(text):
    """Split cleaned PDF text into practical sentence-sized passages."""
    text = re.sub(r"\s+", " ", text or "").strip()

    if not text:
        return []

    return [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+", text)
        if sentence.strip()
    ]


def extract_acronym_expansion_from_text(text, acronym):
    """
    Find a phrase immediately before an acronym in brackets.

    Example:
    Standard variable rate (S V R) -> Standard variable rate
    """
    flexible_acronym = make_spaced_acronym_pattern(acronym)

    pattern = re.compile(
        rf"([A-Za-z][A-Za-z'\-]*(?:\s+[A-Za-z][A-Za-z'\-]*){{0,7}})"
        rf"\s*\(\s*{flexible_acronym}\s*\)",
        flags=re.IGNORECASE,
    )

    best_phrase = ""

    for match in pattern.finditer(text or ""):
        candidate = re.sub(r"\s+", " ", match.group(1)).strip()
        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", candidate)

        # Work backwards to find the shortest suffix whose initials
        # match the acronym. This avoids capturing unrelated words.
        for start in range(len(words)):
            suffix = words[start:]

            if acronym_initials(suffix) == acronym.lower():
                phrase = " ".join(suffix)

                if not best_phrase or len(phrase) < len(best_phrase):
                    best_phrase = phrase

    return best_phrase


def find_exact_acronym_definition(question, chunks):
    """
    Deterministically answer acronym-definition questions before RAG.

    This bypasses embedding thresholds and Gemini when the exact
    definition is visibly present in an uploaded PDF.
    """
    question_lower = (question or "").lower()
    definition_intent = any(
        phrase in question_lower
        for phrase in (
            "definition of",
            "define ",
            "what is",
            "what does",
            "what means",
            "meaning of",
            "stand for",
            "stands for",
        )
    )

    acronyms = extract_acronym_terms(question)

    if not definition_intent or not acronyms:
        return None

    candidates = []

    for chunk in chunks:
        chunk_text = chunk.get("text", "")
        normalised_chunk_terms = set(
            normalise_for_search(chunk_text).split()
        )
        sentences = split_into_sentences(chunk_text)

        for acronym in acronyms:
            if acronym not in normalised_chunk_terms:
                continue

            flexible_pattern = re.compile(
                make_spaced_acronym_pattern(acronym),
                flags=re.IGNORECASE,
            )

            expansion = extract_acronym_expansion_from_text(
                chunk_text,
                acronym,
            )

            matching_sentences = [
                sentence
                for sentence in sentences
                if flexible_pattern.search(sentence)
            ]

            if not matching_sentences:
                continue

            # Prefer a sentence containing a clear definition pattern.
            definition_sentence = next(
                (
                    sentence
                    for sentence in matching_sentences
                    if re.search(
                        r"\b(is|means|refers to|stands for)\b",
                        sentence,
                        flags=re.IGNORECASE,
                    )
                ),
                matching_sentences[0],
            )

            sentence_index = sentences.index(definition_sentence)
            evidence_sentences = [definition_sentence]

            # Add the following sentence when it continues the explanation.
            if sentence_index + 1 < len(sentences):
                next_sentence = sentences[sentence_index + 1]

                if len(next_sentence.split()) <= 35:
                    evidence_sentences.append(next_sentence)

            evidence_text = " ".join(evidence_sentences)
            evidence_text = flexible_pattern.sub(
                acronym.upper(),
                evidence_text,
            )

            # PDF extraction can place a section heading directly before
            # the first definition sentence, producing a duplicate such as:
            # "Standard variable rate (SVR) Standard variable rate (SVR) is..."
            if expansion:
                duplicate_heading_pattern = re.compile(
                    rf"^{re.escape(expansion)}\s*\(\s*{re.escape(acronym.upper())}\s*\)\s*"
                    rf"(?={re.escape(expansion)}\s*\()",
                    flags=re.IGNORECASE,
                )
                evidence_text = duplicate_heading_pattern.sub(
                    "",
                    evidence_text,
                ).strip()

            score = 0

            if expansion:
                score += 5

            if re.search(
                r"\b(is|means|refers to|stands for)\b",
                definition_sentence,
                flags=re.IGNORECASE,
            ):
                score += 4

            if "type of" in definition_sentence.lower():
                score += 3

            candidates.append(
                {
                    "acronym": acronym.upper(),
                    "expansion": expansion,
                    "evidence": evidence_text,
                    "chunk": chunk,
                    "score": score,
                }
            )

    if not candidates:
        return None

    candidates.sort(
        key=lambda item: item["score"],
        reverse=True,
    )
    best = candidates[0]

    expansion_text = (
        f"**{best['acronym']} stands for {best['expansion']}.**\n\n"
        if best["expansion"]
        else ""
    )

    answer = (
        f"{expansion_text}"
        f'The document states: “{best["evidence"]}” '
        f'[Source: {best["chunk"]["source"]}, Page {best["chunk"]["page"]}]'
    )

    selected_chunk = {
        **best["chunk"],
        "score": 2.0,
        "semantic_score": 1.0,
        "keyword_score": 1.0,
        "exact_acronym_match": True,
        "exact_acronym_matches": [best["acronym"].lower()],
        "expansion_match": bool(best["expansion"]),
        "expansion_matches": (
            [
                {
                    "acronym": best["acronym"].lower(),
                    "phrase": best["expansion"],
                }
            ]
            if best["expansion"]
            else []
        ),
    }

    return {
        "question": question,
        "answer": answer,
        "chunks": [selected_chunk],
        "found": True,
        "search_note": (
            "Exact document definition used; similarity threshold bypassed."
        ),
    }


def answer_one_question(question, chunks, embeddings):
    """Search and answer one question using document-grounded evidence."""
    exact_definition = find_exact_acronym_definition(
        question,
        chunks,
    )

    if exact_definition:
        return exact_definition

    relevant_chunks, search_info = retrieve_relevant_chunks(
        question=question,
        chunks=chunks,
        embeddings=embeddings,
    )

    if not relevant_chunks:
        return {
            "question": question,
            "answer": NOT_FOUND_MESSAGE,
            "chunks": [],
            "found": False,
            "search_note": "",
        }

    acronym_expansions = search_info.get("acronym_expansions", {})
    interpretation_lines = []

    for acronym, details in acronym_expansions.items():
        interpretation_lines.append(
            f"{acronym.upper()} = {details['phrase']}"
        )

    interpretation_text = (
        "\n".join(interpretation_lines)
        if interpretation_lines
        else "None"
    )

    context = build_context(relevant_chunks)

    prompt = f"""
You are FinOps AI Copilot supporting a human finance reviewer.

Answer the user's question using only the supplied document evidence.

Rules:
- Do not use outside knowledge.
- Do not invent, infer or assume missing facts.
- Answer clearly and professionally.
- Document-derived acronym interpretations are allowed only when shown below.
- Include inline citations using this exact format:
  [Source: filename.pdf, Page 3]
- Cite every important factual statement.
- If the evidence does not support an answer, reply exactly:
  {NOT_FOUND_MESSAGE}

DOCUMENT-DERIVED ACRONYM INTERPRETATIONS:
{interpretation_text}

DOCUMENT EVIDENCE:
{context}

USER QUESTION:
{question}
"""

    answer = call_gemini(prompt)
    returned_fallback = NOT_FOUND_MESSAGE.lower() in (answer or "").lower()

    focused_chunks = [
        chunk
        for chunk in relevant_chunks
        if chunk.get("expansion_match")
        or chunk.get("exact_acronym_match")
    ]

    # If the first answer misses an acronym definition, retry once with
    # only the document passages that matched the acronym or expansion.
    if returned_fallback and focused_chunks and interpretation_lines:
        focused_context = build_context(focused_chunks)

        retry_prompt = f"""
You are answering an acronym question from uploaded documents.

Use only the evidence and document-derived interpretation below.
Do not use outside knowledge.

DOCUMENT-DERIVED INTERPRETATION:
{interpretation_text}

FOCUSED DOCUMENT EVIDENCE:
{focused_context}

USER QUESTION:
{question}

Explain the term as used in the documents and cite the supporting page.
If the evidence genuinely does not support an answer, reply exactly:
{NOT_FOUND_MESSAGE}
"""

        answer = call_gemini(retry_prompt)
        relevant_chunks = focused_chunks

    search_note = ""

    if interpretation_lines:
        search_note = "Document search interpreted " + "; ".join(
            interpretation_lines
        )

    return {
        "question": question,
        "answer": answer,
        "chunks": relevant_chunks,
        "found": NOT_FOUND_MESSAGE.lower() not in (answer or "").lower(),
        "search_note": search_note,
    }

def answer_question(question, chunks, embeddings, qa_history):
    """Answer one or more questions and preserve them for export."""
    qa_history = list(qa_history or [])

    if not question or not question.strip():
        message = "Please type a question before searching the documents."
        return (
            """
            <div class="answer-box empty-panel">
                <div class="empty-icon">💬</div>
                <strong>Please type a question</strong>
                <span>Ask about facts, risks, policy rules, exceptions or evidence in the uploaded documents.</span>
            </div>
            """,
            qa_history,
            message,
            "",
        )

    if not chunks or embeddings is None:
        message = "Process documents before asking a question."
        return (
            """
            <div class="answer-box empty-panel">
                <div class="empty-icon">📂</div>
                <strong>Process finance documents first</strong>
                <span>Upload the expense, receipt and policy documents and select “Process Finance Documents”.</span>
            </div>
            """,
            qa_history,
            message,
            "",
        )

    subquestions = split_into_subquestions(question)
    results = [
        answer_one_question(
            subquestion,
            chunks,
            embeddings,
        )
        for subquestion in subquestions
    ]

    multi_question_notice = ""

    if len(subquestions) > 1:
        multi_question_notice = f"""
        <div class="question-detection">
            ✅ Detected and searched {len(subquestions)} separate questions.
        </div>
        """

    answer_sections = []
    copy_sections = []

    for index, result in enumerate(results, start=1):
        safe_question = html.escape(result["question"])
        safe_answer = format_ai_text(result["answer"])

        number_label = (
            f"<span class='answer-number'>{index}</span>"
            if len(results) > 1
            else ""
        )

        sources_html = format_source_list(result["chunks"])
        search_note_html = ""

        if result.get("search_note"):
            search_note_html = f"""
            <div class="search-interpretation">
                🔤 {html.escape(result['search_note'])}
            </div>
            """

        answer_sections.append(
            f"""
            <div class="individual-answer">
                <div class="question-heading">
                    {number_label}
                    <span>{safe_question}</span>
                </div>
                {search_note_html}
                <div class="answer-text formatted-ai-text">{safe_answer}</div>
                {sources_html}
            </div>
            """
        )

        unique_sources = []
        seen_sources = set()
        for chunk in result["chunks"]:
            source_pair = (chunk["source"], int(chunk["page"]))
            if source_pair not in seen_sources:
                seen_sources.add(source_pair)
                unique_sources.append(
                    {"source": source_pair[0], "page": source_pair[1]}
                )

        qa_entry = {
            "question": result["question"],
            "answer": result["answer"],
            "sources": unique_sources,
            "search_note": result.get("search_note", ""),
        }
        qa_history.append(qa_entry)

        copy_lines = [
            f"Question {index}: {result['question']}",
            "",
            result["answer"],
        ]
        if unique_sources:
            copy_lines.extend(["", "Supporting sources:"])
            copy_lines.extend(
                f"- {item['source']} — Page {item['page']}"
                for item in unique_sources
            )
        copy_sections.append("\n".join(copy_lines))

    answer_html = f"""
    <div class="answer-box">
        <div class="panel-heading">
            <span>💬 Document-grounded answer</span>
        </div>
        {multi_question_notice}
        {''.join(answer_sections)}
    </div>
    """

    copy_text = "DOCUMENT-GROUNDED ANSWER\n\n" + "\n\n---\n\n".join(copy_sections)
    return answer_html, qa_history, copy_text, ""




# ---------------------------------------------------------
# Copy and export helpers
# ---------------------------------------------------------

def safe_export_name(value):
    """Create a Windows- and web-safe export filename fragment."""
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip())
    return cleaned.strip("._-") or "document_review"


def review_payload_to_copy_text(review_payload):
    """Convert a structured review into clean tab-separated text."""
    if not review_payload:
        return "No structured review has been generated yet."

    review_key = review_payload.get("review_key", "")
    config = STRUCTURED_REVIEW_CONFIG.get(review_key)

    if not config:
        return clean_text(review_payload.get("raw_answer", "")) or NOT_FOUND_MESSAGE

    lines = [
        f"FINANCE REVIEW — {review_payload.get('title', config['title']).upper()}",
        "",
        f"Review summary: {review_payload.get('summary', NOT_FOUND_MESSAGE)}",
    ]

    items = review_payload.get("items", [])
    if not items:
        return "\n".join(lines)

    lines.extend(["", "\t".join(label for label, _ in config["columns"])])
    for item in items:
        lines.append(
            "\t".join(
                clean_text(str(item.get(key, "Not stated"))) or "Not stated"
                for _, key in config["columns"]
            )
        )

    return "\n".join(lines)


def generate_current_review_csv(current_review_key, review_results):
    """Create a CSV download for the currently displayed finance review."""
    review_results = review_results or {}
    review_payload = review_results.get(current_review_key or "")

    if not review_payload or not current_review_key:
        return None

    config = STRUCTURED_REVIEW_CONFIG.get(current_review_key)
    if not config:
        return None

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    title_part = safe_export_name(config["title"])
    csv_path = EXPORT_DIR / f"FinOps_AI_Copilot_{title_part}_{timestamp}.csv"

    with csv_path.open("w", newline="", encoding="utf-8-sig") as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["FinOps Review", config["title"]])
        writer.writerow(["Summary", review_payload.get("summary", NOT_FOUND_MESSAGE)])
        writer.writerow([])
        writer.writerow([label for label, _ in config["columns"]])

        for item in review_payload.get("items", []):
            writer.writerow(
                [item.get(key, "Not stated") for _, key in config["columns"]]
            )

    return str(csv_path)


def add_markdownish_text(document, text):
    """Add lightweight Markdown-style text to a Word document."""
    lines = (text or "").replace("\r\n", "\n").split("\n")
    added_any = False

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        line = re.sub(r"__(.*?)__", r"\1", line)

        if re.match(r"^#{1,4}\s+", line):
            heading_text = re.sub(r"^#{1,4}\s+", "", line)
            document.add_heading(heading_text, level=2)
        elif re.match(r"^(?:[-*•]|\d+[.)])\s+", line):
            bullet_text = re.sub(r"^(?:[-*•]|\d+[.)])\s+", "", line)
            document.add_paragraph(bullet_text, style="List Bullet")
        else:
            document.add_paragraph(line)
        added_any = True

    if not added_any:
        document.add_paragraph("No content was generated for this section.")


def set_docx_defaults(document):
    """Apply a clean business-report layout."""
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"


def add_review_table_to_docx(document, review_key, payload):
    """Add one structured review section to the Word report."""
    config = STRUCTURED_REVIEW_CONFIG[review_key]
    document.add_heading(config["title"], level=1)
    document.add_paragraph(payload.get("summary", NOT_FOUND_MESSAGE))

    items = payload.get("items", [])
    if not items:
        document.add_paragraph("No structured rows were found for this review.")
        return

    report_columns = config["columns"]

    table = document.add_table(rows=1, cols=len(report_columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, (label, _) in enumerate(report_columns):
        cell = table.rows[0].cells[index]
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for item in items:
        cells = table.add_row().cells
        for index, (_, key) in enumerate(report_columns):
            value = str(item.get(key, "Not stated"))

            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    document.add_paragraph()


def generate_missing_reviews_for_report(chunks, review_results):
    """Generate any missing finance review sections in one Gemini request."""
    updated_results = dict(review_results or {})
    missing_keys = [
        key
        for key in STRUCTURED_REVIEW_CONFIG
        if key not in updated_results
        or not isinstance(updated_results.get(key), dict)
        or updated_results[key].get("format") == "narrative"
    ]

    if not missing_keys:
        return updated_results, None

    context, _ = build_balanced_review_context(chunks, max_characters=45000)
    if not context:
        return updated_results, "No readable document evidence was available."

    combined_schema = {
        key: STRUCTURED_REVIEW_CONFIG[key]["schema"]
        for key in missing_keys
    }
    instructions = "\n".join(
        f"- {key}: {COMPLETE_REVIEW_INSTRUCTIONS[key]}"
        for key in missing_keys
    )

    prompt = f"""
You are FinOps AI Copilot supporting a human finance reviewer.

Create the missing structured finance sections for a complete review report using
only the supplied document evidence.

SECTIONS TO CREATE:
{instructions}

Rules:
- Return ONLY valid JSON. Do not use Markdown or code fences.
- The top-level JSON keys must exactly match the section keys in the schema.
- Use the exact nested keys shown in each section schema.
- Use string values for every field.
- Use exact uploaded file names and actual page numbers from the evidence.
- Do not use outside knowledge or invent missing facts.
- If policy explicitly requires prior or advance approval and the evidence says it was not obtained, do not recommend creating or obtaining that prior approval after the event unless the policy explicitly permits retrospective approval.
- Never suggest backdating or fabricating approval evidence.
- If a field is not stated, write "Not stated".
- If a section has no relevant information, return an empty items list and
  put "{NOT_FOUND_MESSAGE}" in its summary.

EXACT COMBINED JSON SCHEMA:
{json.dumps(combined_schema, indent=2, ensure_ascii=False)}

DOCUMENT EVIDENCE:
{context}
"""

    answer = call_gemini(prompt)
    if answer == GEMINI_BUSY_MESSAGE:
        return updated_results, GEMINI_BUSY_MESSAGE

    parsed = extract_json_object(answer)
    if not parsed:
        return updated_results, (
            "Gemini did not return the structured report format. "
            "Please try generating the report again."
        )

    for key in missing_keys:
        section_payload = parsed.get(key, {})
        if isinstance(section_payload, dict):
            updated_results[key] = normalise_review_payload(section_payload, key)

    still_missing = [key for key in missing_keys if key not in updated_results]
    if still_missing:
        return updated_results, (
            "The report could not create every review section. "
            "Please try again."
        )

    return updated_results, None


def build_word_report(summary_text, processed_documents, review_results, qa_history, ai_recommendation="", hitl_decision="", reviewer_note=""):
    """Create a complete editable FinOps finance review report."""
    timestamp = datetime.now()
    report_path = EXPORT_DIR / (
        f"FinOps_AI_Copilot_Report_{timestamp.strftime('%Y%m%d_%H%M%S')}.docx"
    )

    document = Document()
    set_docx_defaults(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("FinOps AI Copilot — Finance Review Report")
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"Generated {timestamp.strftime('%d %B %Y at %H:%M')}"
    ).italic = True

    document.add_heading("Finance Documents Reviewed", level=1)
    if processed_documents:
        for item in processed_documents:
            document.add_paragraph(
                (
                    f"{item.get('source', 'Unknown document')} — "
                    f"{item.get('readable_pages', 0)} readable / "
                    f"{item.get('total_pages', 0)} total pages"
                    f" — {item.get('extraction_method', 'Text extraction')}"
                ),
                style="List Bullet",
            )
    else:
        document.add_paragraph("No document metadata was available.")

    document.add_heading("Finance Case Summary", level=1)
    add_markdownish_text(document, summary_text)

    for review_key in ("risks", "actions", "dates", "comparison"):
        payload = (review_results or {}).get(review_key)
        if payload:
            add_review_table_to_docx(document, review_key, payload)
        else:
            document.add_heading(
                STRUCTURED_REVIEW_CONFIG[review_key]["title"],
                level=1,
            )
            document.add_paragraph("This section was not generated.")

    document.add_heading("Reviewer Questions and Answers", level=1)
    if qa_history:
        for index, entry in enumerate(qa_history, start=1):
            document.add_heading(
                f"Question {index}: {entry.get('question', '')}",
                level=2,
            )
            add_markdownish_text(document, entry.get("answer", ""))

            sources = entry.get("sources", [])
            if sources:
                source_paragraph = document.add_paragraph()
                source_paragraph.add_run("Supporting sources").bold = True
                for source in sources:
                    document.add_paragraph(
                        f"{source.get('source', 'Unknown')} — Page {source.get('page', 'Not stated')}",
                        style="List Bullet",
                    )
    else:
        document.add_paragraph(
            "No finance or policy questions were asked during this session."
        )

    document.add_heading("AI Recommendation (Advisory)", level=1)
    if ai_recommendation:
        add_markdownish_text(document, ai_recommendation)
    else:
        document.add_paragraph("No AI recommendation was generated before export.")

    document.add_heading("Human-in-the-Loop Finance Decision", level=1)
    document.add_paragraph(f"Final reviewer decision: {hitl_decision or 'Not recorded'}")
    document.add_paragraph(f"Reviewer note: {reviewer_note or 'No reviewer note entered.'}")

    document.add_heading("Important Notice", level=1)
    document.add_paragraph(
        "This report was generated with AI from the uploaded finance documents. "
        "AI can make mistakes. The AI recommendation is advisory only. "
        "A human finance reviewer retains the final Approve, Reject or Escalate decision."
    )

    document.save(report_path)
    return str(report_path)


def generate_complete_report(
    chunks,
    embeddings,
    summary_text,
    processed_documents,
    review_results,
    qa_history,
    ai_recommendation,
    hitl_decision,
    reviewer_note,
):
    """Generate missing finance reviews and package the session into Word."""
    if not chunks or embeddings is None:
        return (
            None,
            """
            <div class="status-box error">
                <strong>Process finance documents first.</strong><br>
                Upload and process the finance documents before creating a report.
            </div>
            """,
            review_results or {},
        )

    if not summary_text:
        return (
            None,
            """
            <div class="status-box error">
                <strong>No summary is available.</strong><br>
                Process the documents again before creating a report.
            </div>
            """,
            review_results or {},
        )

    try:
        complete_results, generation_error = generate_missing_reviews_for_report(
            chunks,
            review_results,
        )

        if generation_error:
            return (
                None,
                f"""
                <div class="status-box error">
                    <strong>The complete report could not be generated yet.</strong><br>
                    {html.escape(generation_error)}
                </div>
                """,
                complete_results,
            )

        report_path = build_word_report(
            summary_text,
            processed_documents,
            complete_results,
            qa_history,
            ai_recommendation,
            hitl_decision,
            reviewer_note,
        )

        return (
            report_path,
            """
            <div class="status-box success export-success">
                <strong>✅ Word report ready</strong><br>
                The report includes the finance case summary, all four structured
                finance reviews, AI recommendation, HITL decision, document references
                and any reviewer questions answered in this session. Use the download control below.
            </div>
            """,
            complete_results,
        )

    except Exception as error:
        return (
            None,
            f"""
            <div class="status-box error">
                <strong>Report generation error:</strong><br>
                {html.escape(str(error))}
            </div>
            """,
            review_results or {},
        )


# ---------------------------------------------------------
# Quick Review buttons
# ---------------------------------------------------------

def build_balanced_review_context(chunks, max_characters=40000):
    """
    Build broad evidence across all uploaded documents.

    Finance Quick Review deliberately does not use the normal question splitter
    or minimum similarity threshold because reviews such as risks,
    obligations and comparisons require wider document coverage.
    """
    if not chunks:
        return "", []

    chunks_by_source = {}

    for chunk in chunks:
        chunks_by_source.setdefault(chunk["source"], []).append(chunk)

    selected_chunks = []
    context_sections = []
    document_count = max(len(chunks_by_source), 1)
    per_document_limit = max(5000, max_characters // document_count)

    for source_name, source_chunks in chunks_by_source.items():
        source_chunks = sorted(
            source_chunks,
            key=lambda item: (
                item.get("page", 0),
                item.get("chunk_number", 0),
            ),
        )

        current_length = 0

        for chunk in source_chunks:
            section = (
                f"Source: {chunk['source']}\n"
                f"Page: {chunk['page']}\n"
                f"Text: {chunk['text']}"
            )

            remaining = per_document_limit - current_length

            if remaining <= 0:
                break

            if len(section) > remaining:
                section = section[:remaining]

            context_sections.append(section)
            selected_chunks.append(chunk)
            current_length += len(section)

    return "\n\n---\n\n".join(context_sections), selected_chunks


def find_chunks_cited_in_answer(answer, available_chunks):
    """Match Gemini inline citations back to source/page chunks."""
    citation_pattern = re.compile(
        r"\[Source:\s*(.+?),\s*Page\s*(\d+)\]",
        flags=re.IGNORECASE,
    )

    cited_pairs = []

    for source_name, page_number in citation_pattern.findall(answer or ""):
        pair = (source_name.strip(), int(page_number))

        if pair not in cited_pairs:
            cited_pairs.append(pair)

    cited_chunks = []

    for source_name, page_number in cited_pairs:
        matching_chunk = next(
            (
                chunk
                for chunk in available_chunks
                if chunk["source"].lower() == source_name.lower()
                and int(chunk["page"]) == page_number
            ),
            None,
        )

        if matching_chunk:
            cited_chunks.append(matching_chunk)

    return cited_chunks


def extract_json_object(text):
    """Extract the first valid JSON object from a Gemini response."""
    cleaned = (text or "").strip()

    if not cleaned:
        return None

    cleaned = re.sub(
        r"^```(?:json)?\s*|\s*```$",
        "",
        cleaned,
        flags=re.IGNORECASE | re.DOTALL,
    ).strip()

    first_brace = cleaned.find("{")
    last_brace = cleaned.rfind("}")

    if first_brace == -1 or last_brace == -1 or last_brace <= first_brace:
        return None

    candidate = cleaned[first_brace:last_brace + 1]

    try:
        parsed = json.loads(candidate)
    except json.JSONDecodeError:
        return None

    return parsed if isinstance(parsed, dict) else None


def normalise_review_payload(payload, review_key):
    """Validate and normalise one structured finance review payload."""
    config = STRUCTURED_REVIEW_CONFIG[review_key]
    valid_keys = [key for _, key in config["columns"]]

    summary = clean_text(str((payload or {}).get("summary", "")))
    raw_items = (payload or {}).get("items", [])

    if not isinstance(raw_items, list):
        raw_items = []

    normalised_items = []

    for raw_item in raw_items:
        if not isinstance(raw_item, dict):
            continue

        item = {}

        for key in valid_keys:
            value = raw_item.get(key, "")
            value = clean_text(str(value)) if value is not None else ""

            if key.startswith("page") and value:
                page_match = re.search(r"\d+", value)
                value = page_match.group(0) if page_match else value

            item[key] = value or "Not stated"

        normalised_items.append(item)

    return {
        "review_key": review_key,
        "title": config["title"],
        "summary": summary or NOT_FOUND_MESSAGE,
        "items": normalised_items,
    }


def find_chunks_from_review_items(review_payload, available_chunks):
    """Match structured source/page fields back to document chunks."""
    source_page_pairs = []

    for item in review_payload.get("items", []):
        for source_key, page_key in (
            ("source", "page"),
            ("source_a", "page_a"),
            ("source_b", "page_b"),
        ):
            source_name = clean_text(item.get(source_key, ""))
            page_text = clean_text(item.get(page_key, ""))

            if not source_name or source_name == "Not stated":
                continue

            page_match = re.search(r"\d+", page_text)

            if not page_match:
                continue

            pair = (source_name, int(page_match.group(0)))

            if pair not in source_page_pairs:
                source_page_pairs.append(pair)

    matched_chunks = []

    for source_name, page_number in source_page_pairs:
        matching_chunk = next(
            (
                chunk
                for chunk in available_chunks
                if chunk["source"].lower() == source_name.lower()
                and int(chunk["page"]) == page_number
            ),
            None,
        )

        if matching_chunk:
            matched_chunks.append(matching_chunk)

    return matched_chunks


def make_structured_review_html(review_payload, evidence_chunks):
    """Render a structured Quick Review as a clear HTML table."""
    review_key = review_payload["review_key"]
    config = STRUCTURED_REVIEW_CONFIG[review_key]
    title = review_payload["title"]
    summary = review_payload.get("summary", "")
    items = review_payload.get("items", [])

    if not items:
        table_html = f"""
        <div class="structured-empty">
            {html.escape(summary or NOT_FOUND_MESSAGE)}
        </div>
        """
    else:
        header_html = "".join(
            f"<th>{html.escape(label)}</th>"
            for label, _ in config["columns"]
        )

        rows = []

        for item in items:
            cells = []

            for _, key in config["columns"]:
                value = html.escape(str(item.get(key, "Not stated")))

                if key == "severity":
                    severity_class = re.sub(
                        r"[^a-z]+",
                        "-",
                        value.lower(),
                    ).strip("-")
                    value = (
                        f"<span class='severity-badge severity-{severity_class}'>"
                        f"{value}</span>"
                    )

                cells.append(f"<td>{value}</td>")

            rows.append(f"<tr>{''.join(cells)}</tr>")

        table_html = f"""
        <div class="structured-table-wrap">
            <table class="structured-review-table">
                <thead><tr>{header_html}</tr></thead>
                <tbody>{''.join(rows)}</tbody>
            </table>
        </div>
        """

    cited_chunks = find_chunks_from_review_items(
        review_payload,
        evidence_chunks,
    )

    if not cited_chunks:
        cited_chunks = evidence_chunks[:12]

    sources_html = format_source_list(cited_chunks)

    return f"""
    <div class="answer-box">
        <div class="panel-heading">
            <span>⚡ Finance Quick Review — {html.escape(title)}</span>
        </div>
        <div class="quick-review-success">
            ✅ Completed as one structured document-wide review.
        </div>
        <div class="structured-review-summary">
            <strong>Review summary</strong>
            <span>{html.escape(summary)}</span>
        </div>
        {table_html}
        {sources_html}
    </div>
    """


def make_quick_review_html(review_title, answer, evidence_chunks):
    """Display a narrative fallback if structured JSON cannot be parsed."""
    if answer == GEMINI_BUSY_MESSAGE:
        return f"""
        <div class="answer-box">
            <div class="panel-heading">
                <span>⚡ Finance Quick Review — {html.escape(review_title)}</span>
            </div>
            <div class="service-warning">
                <strong>⏳ Gemini is temporarily busy</strong>
                <span>{html.escape(GEMINI_BUSY_MESSAGE)}</span>
            </div>
        </div>
        """

    formatted_answer = format_ai_text(answer)
    cited_chunks = find_chunks_cited_in_answer(
        answer,
        evidence_chunks,
    )

    if not cited_chunks:
        cited_chunks = evidence_chunks[:12]

    sources_html = format_source_list(cited_chunks)

    return f"""
    <div class="answer-box">
        <div class="panel-heading">
            <span>⚡ Finance Quick Review — {html.escape(review_title)}</span>
        </div>
        <div class="quick-review-success">
            ✅ Completed as one document-wide review.
        </div>
        <div class="structured-fallback-note">
            The review was completed, but Gemini returned narrative text
            instead of the requested table format.
        </div>
        <div class="answer-text formatted-ai-text">
            {formatted_answer}
        </div>
        {sources_html}
    </div>
    """


def run_quick_review(
    review_key,
    display_question,
    review_instruction,
    chunks,
    embeddings,
    review_results,
):
    """Run one document-wide review and preserve export-ready results."""
    review_results = dict(review_results or {})
    config = STRUCTURED_REVIEW_CONFIG[review_key]
    review_title = config["title"]

    if not chunks or embeddings is None:
        message = "Process documents before running a finance review."
        return (
            display_question,
            """
            <div class="answer-box empty-panel">
                <div class="empty-icon">📂</div>
                <strong>Process finance documents first</strong>
                <span>Upload the expense, receipt and policy documents and select “Process Finance Documents”.</span>
            </div>
            """,
            review_results,
            message,
            "",
        )

    context, evidence_chunks = build_balanced_review_context(chunks)

    if not context:
        message = "No readable evidence was found in the uploaded documents."
        return (
            display_question,
            """
            <div class="answer-box empty-panel">
                <div class="empty-icon">⚠️</div>
                <strong>No readable evidence</strong>
                <span>The uploaded documents did not contain readable text.</span>
            </div>
            """,
            review_results,
            message,
            "",
        )

    schema_text = json.dumps(
        config["schema"],
        indent=2,
        ensure_ascii=False,
    )

    prompt = f"""
You are FinOps AI Copilot supporting a human finance reviewer.

Complete one finance document-wide review using only the supplied evidence.

REVIEW TYPE:
{review_title}

REVIEW INSTRUCTION:
{review_instruction}

Rules:
- Treat this as one review task, not as separate questions.
- Review evidence across all uploaded documents.
- Do not use outside knowledge.
- Do not invent missing facts.
- Return ONLY valid JSON. Do not use Markdown or code fences.
- Use the exact keys shown in the JSON schema below.
- Use string values for every field.
- Use the exact uploaded file name shown in the evidence.
- Use the actual page number shown in the evidence.
- Add one row per distinct finding.
- If a field is not stated, write "Not stated".
- If policy explicitly requires prior or advance approval and the evidence says it was not obtained, do not suggest creating, backdating or obtaining that prior approval after the event. Treat it as an exception and recommend deduction/rejection or human escalation as supported by the policy.
- Do not describe already-missing prior approval as evidence that can simply be supplied later unless the documents say such retrospective approval is permitted.
- If no relevant information exists, return an empty items list and put
  "{NOT_FOUND_MESSAGE}" in summary.

EXACT JSON SCHEMA:
{schema_text}

DOCUMENT EVIDENCE:
{context}
"""

    answer = call_gemini(prompt)

    if answer == GEMINI_BUSY_MESSAGE:
        return (
            display_question,
            make_quick_review_html(
                review_title,
                answer,
                evidence_chunks,
            ),
            review_results,
            GEMINI_BUSY_MESSAGE,
            "",
        )

    parsed_payload = extract_json_object(answer)

    if parsed_payload is not None:
        review_payload = normalise_review_payload(
            parsed_payload,
            review_key,
        )
        review_results[review_key] = review_payload

        return (
            display_question,
            make_structured_review_html(
                review_payload,
                evidence_chunks,
            ),
            review_results,
            review_payload_to_copy_text(review_payload),
            review_key,
        )

    review_payload = {
        "review_key": review_key,
        "title": review_title,
        "summary": "",
        "items": [],
        "raw_answer": answer,
        "format": "narrative",
    }
    review_results[review_key] = review_payload

    return (
        display_question,
        make_quick_review_html(
            review_title,
            answer,
            evidence_chunks,
        ),
        review_results,
        answer or NOT_FOUND_MESSAGE,
        "",
    )



def review_key_risks(chunks, embeddings, review_results):
    return run_quick_review(
        "risks",
        "Finance Review: Expense Summary",
        COMPLETE_REVIEW_INSTRUCTIONS["risks"],
        chunks,
        embeddings,
        review_results,
    )
def review_actions_obligations(chunks, embeddings, review_results):
    return run_quick_review(
        "actions",
        "Finance Review: Policy Compliance",
        COMPLETE_REVIEW_INSTRUCTIONS["actions"],
        chunks,
        embeddings,
        review_results,
    )
def review_important_dates(chunks, embeddings, review_results):
    return run_quick_review(
        "dates",
        "Finance Review: Exceptions & Risks",
        COMPLETE_REVIEW_INSTRUCTIONS["dates"],
        chunks,
        embeddings,
        review_results,
    )
def review_document_comparison(chunks, embeddings, review_results):
    return run_quick_review(
        "comparison",
        "Finance Review: Missing Evidence",
        COMPLETE_REVIEW_INSTRUCTIONS["comparison"],
        chunks,
        embeddings,
        review_results,
    )
# ---------------------------------------------------------
# AI recommendation + Human-in-the-Loop finance decision
# ---------------------------------------------------------

def generate_finance_recommendation(chunks, review_results):
    """Generate an advisory AI recommendation while preserving human final control."""
    if not chunks:
        return (
            """
            <div class="answer-box empty-panel">
                <div class="empty-icon">📂</div>
                <strong>Process finance documents first</strong>
                <span>Upload the expense, receipt and policy documents before generating a recommendation.</span>
            </div>
            """,
            "",
            review_results or {},
        )

    complete_results, generation_error = generate_missing_reviews_for_report(
        chunks,
        review_results or {},
    )

    if generation_error:
        return (
            f"""
            <div class="answer-box service-warning-panel">
                <div class="service-warning">
                    <strong>AI recommendation is not ready</strong>
                    <span>{html.escape(generation_error)}</span>
                </div>
            </div>
            """,
            "",
            complete_results,
        )

    context, _ = build_balanced_review_context(chunks, max_characters=30000)
    review_json = json.dumps(complete_results, indent=2, ensure_ascii=False)

    prompt = f"""
You are FinOps AI Copilot supporting a human finance reviewer.

Using only the finance review results and uploaded document evidence below,
provide an ADVISORY recommendation. The human reviewer makes the final decision.

Choose exactly one recommendation:
- APPROVE
- REJECT
- ESCALATE / MANUAL REVIEW

Return a concise response in this format:
Recommendation: <one option above>
Confidence: <High, Medium or Low>
Rationale:
- <reason 1>
- <reason 2>
- <reason 3 if useful>
Required follow-up: <what Finance should check or request next, or None identified>

Rules:
- Do not use outside knowledge.
- Do not invent policy limits, receipts, approvals or evidence.
- If material evidence is missing or policy interpretation is uncertain, prefer ESCALATE / MANUAL REVIEW.
- If a rule requires prior or advance approval and the documents state that approval was not obtained, do not tell Finance to obtain or provide that approval retrospectively unless the uploaded policy explicitly permits retrospective approval.
- Never suggest backdating or fabricating evidence. Where prior approval is absent, recommend the policy-supported consequence (for example deduction/rejection) or human escalation.
- This is an AI recommendation only, not the final finance decision.

STRUCTURED FINANCE REVIEW RESULTS:
{review_json}

DOCUMENT EVIDENCE:
{context}
"""

    recommendation = call_gemini(prompt)

    if recommendation == GEMINI_BUSY_MESSAGE:
        return (
            make_quick_review_html("AI Recommendation", recommendation, []),
            "",
            complete_results,
        )

    recommendation = recommendation or "No recommendation was generated."
    recommendation_html = f"""
    <div class="answer-box">
        <div class="panel-heading"><span>🤖 AI Recommendation — Advisory Only</span></div>
        <div class="safety-note">
            <strong>Human-in-the-Loop:</strong> Finance retains the final decision.
        </div>
        <div class="answer-text formatted-ai-text">
            {format_ai_text(recommendation)}
        </div>
    </div>
    """

    return recommendation_html, recommendation, complete_results


def record_hitl_decision(decision, reviewer_note, ai_recommendation):
    """Record the human Finance decision for the current session/report."""
    note = clean_text(reviewer_note or "")
    ai_ready = bool(clean_text(ai_recommendation or ""))
    decision_text = decision

    note_html = html.escape(note) if note else "No reviewer note entered."
    recommendation_note = (
        "AI recommendation was generated before this decision."
        if ai_ready
        else "No AI recommendation was generated before this decision."
    )

    status_html = f"""
    <div class="status-box success">
        <div class="status-title">👤 Human Finance Decision: {html.escape(decision)}</div>
        <div class="helper-text">{html.escape(recommendation_note)}</div>
        <div class="helper-text"><strong>Reviewer note:</strong> {note_html}</div>
    </div>
    """
    return status_html, decision_text


def hitl_approve(reviewer_note, ai_recommendation):
    return record_hitl_decision("APPROVED", reviewer_note, ai_recommendation)


def hitl_reject(reviewer_note, ai_recommendation):
    return record_hitl_decision("REJECTED", reviewer_note, ai_recommendation)


def hitl_escalate(reviewer_note, ai_recommendation):
    return record_hitl_decision("ESCALATED / MANUAL REVIEW", reviewer_note, ai_recommendation)


def clear_all():
    """Reset all visible fields and temporary finance-review memory."""
    return (
        None,                       # pdf_input
        "",                         # status_output
        EMPTY_SUMMARY_HTML,         # summary_output
        "",                         # question_input
        EMPTY_ANSWER_HTML,          # answer_output
        [],                         # chunks_state
        None,                       # embeddings_state
        {},                         # review_results_state
        "",                         # summary_text_state
        [],                         # processed_documents_state
        [],                         # qa_history_state
        "",                         # current_copy_text_state
        "",                         # current_review_key_state
        "",                         # summary_copy_status
        "",                         # result_copy_status
        None,                       # report_file_output
        "",                         # report_status_output
        EMPTY_RECOMMENDATION_HTML,  # ai_recommendation_output
        "",                         # ai_recommendation_state
        "",                         # hitl_decision_state
        "",                         # reviewer_note
        "",                         # hitl_status_output
    )


# ---------------------------------------------------------
# Interface styling
# ---------------------------------------------------------

css = """
:root {
    --primary: #5850ec;
    --primary-dark: #4338ca;
    --primary-soft: #eef2ff;
    --border: #dbe2f2;
    --text: #172033;
    --muted: #64748b;
    --surface: #ffffff;
    --background: #f5f7fc;
    --finops-header-height: 118px;
}

html,
body {
    width: 100% !important;
    height: 100% !important;
    min-height: 0 !important;
    margin: 0 !important;
    overflow: hidden !important;
    background: var(--background) !important;
}

gradio-app {
    display: block !important;
    width: 100% !important;
    height: 100vh !important;
    height: 100dvh !important;
    min-height: 0 !important;
    overflow: hidden !important;
}

.gradio-container {
    max-width: none !important;
    width: 100% !important;
    height: 100vh !important;
    height: 100dvh !important;
    min-height: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: hidden !important;
    box-sizing: border-box !important;
}

/* v0.8: ONE authoritative scroll container.
   Header + Steps 1-6 are children of this same Gradio Column. */
#finops-scroll-shell {
    width: 100% !important;
    height: 100vh !important;
    height: 100dvh !important;
    max-height: 100vh !important;
    max-height: 100dvh !important;
    min-height: 0 !important;
    margin: 0 !important;
    padding: 0 12px 80px 12px !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    position: relative !important;
    box-sizing: border-box !important;
    scrollbar-gutter: stable;
    overscroll-behavior: contain;
    -webkit-overflow-scrolling: touch;
    gap: 0 !important;
}

/* Primary option: the WHOLE Gradio HTML component is sticky. */
#finops-topbar {
    position: sticky !important;
    position: -webkit-sticky !important;
    top: 0 !important;
    z-index: 10000 !important;
    width: 100% !important;
    margin: 0 !important;
    padding: 0 !important;
    border: 0 !important;
    background: rgba(248, 250, 255, 0.995) !important;
    box-shadow: 0 8px 20px rgba(19, 33, 60, 0.10) !important;
    box-sizing: border-box !important;
}

#finops-topbar,
#finops-topbar > div,
#finops-topbar .prose {
    overflow: visible !important;
    max-height: none !important;
}

#finops-topbar-spacer {
    display: none !important;
    width: 100% !important;
    pointer-events: none !important;
}

/* Automatic fallback option:
   if the browser/Gradio wrapper defeats sticky, JS applies this class.
   The header then becomes fixed to the app viewport and a spacer preserves layout. */
#finops-scroll-shell.finops-fixed-fallback #finops-topbar {
    position: fixed !important;
    top: 0 !important;
    left: var(--finops-shell-left, 0px) !important;
    width: var(--finops-shell-width, 100%) !important;
    right: auto !important;
    z-index: 20000 !important;
}

#finops-scroll-shell.finops-fixed-fallback #finops-topbar-spacer {
    display: block !important;
    height: var(--finops-topbar-height, 108px) !important;
    min-height: var(--finops-topbar-height, 108px) !important;
}

/* Make the top box deliberately compact so more workspace is visible. */
.app-hero {
    position: static;
    text-align: center;
    margin: 0;
    padding: 5px 14px 7px 14px;
    background: rgba(248, 250, 255, 0.995);
    border-bottom: 1px solid rgba(88, 80, 236, 0.14);
    backdrop-filter: blur(8px);
    -webkit-backdrop-filter: blur(8px);
}

/* Keep Gradio progress/status overlays inside cards. */
.section-card,
.section-card > div,
.section-card .wrap {
    min-width: 0 !important;
    box-sizing: border-box !important;
}

.section-card {
    overflow-x: hidden !important;
}

.section-card [class*="progress"],
.section-card [class*="eta"],
.section-card [data-testid*="progress"],
.section-card [data-testid*="status"] {
    max-width: calc(100% - 12px) !important;
    right: 6px !important;
    left: auto !important;
    box-sizing: border-box !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    white-space: nowrap !important;
}


/* v0.9: show Gradio processing feedback again, but keep the
   "processing | 12.3s" runtime text inside the component boundary. */
[data-testid="status-tracker"] .progress-text.meta-text,
[data-testid="status-tracker"] .progress-text {
    box-sizing: border-box !important;
    max-width: calc(100% - 12px) !important;
    right: 6px !important;
    left: auto !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    font-size: 11px !important;
    line-height: 1.2 !important;
    padding: 3px 5px !important;
    border-radius: 5px !important;
}

/* Keep the status tracker clipped to the component instead of allowing
   the timer text to protrude outside narrow result boxes. */
[data-testid="status-tracker"] {
    max-width: 100% !important;
    overflow: hidden !important;
    border-radius: inherit !important;
}



/* v0.10: clear, app-owned processing feedback.
   We no longer rely on Gradio's tiny runtime text in the top-right corner. */
.finops-processing {
    width: 100%;
    min-height: 86px;
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 12px;
    padding: 18px 20px;
    box-sizing: border-box;
    border: 1px solid rgba(88, 80, 236, 0.28);
    border-radius: 12px;
    background: linear-gradient(180deg, rgba(246,247,255,0.98), rgba(250,250,255,0.98));
    color: #17213a;
    text-align: left;
}

.finops-processing-spinner {
    width: 22px;
    height: 22px;
    min-width: 22px;
    border: 3px solid rgba(88, 80, 236, 0.18);
    border-top-color: #5b55ed;
    border-radius: 50%;
    animation: finops-spin 0.8s linear infinite;
}

.finops-processing-copy {
    display: flex;
    flex-direction: column;
    gap: 2px;
}

.finops-processing-copy strong {
    font-size: 14px;
    line-height: 1.3;
    color: #17213a;
}

.finops-processing-copy span {
    font-size: 12px;
    line-height: 1.35;
    color: #667085;
}

@keyframes finops-spin {
    to { transform: rotate(360deg); }
}


.title-line {
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 10px;
    flex-wrap: wrap;
}

.main-title {
    font-size: 28px;
    line-height: 1.08;
    font-weight: 850;
    color: #13213c;
}

.version-badge {
    display: inline-flex;
    align-items: center;
    padding: 5px 11px;
    border-radius: 999px;
    background: linear-gradient(135deg, #5850ec, #7c3aed);
    color: white;
    font-size: 13px;
    font-weight: 800;
    letter-spacing: 0.3px;
    box-shadow: 0 5px 14px rgba(88, 80, 236, 0.25);
}

.sub-title {
    margin: 3px auto 0 auto;
    max-width: 1100px;
    font-size: 12.5px;
    line-height: 1.35;
    color: var(--muted);
}

.feature-strip {
    display: flex;
    justify-content: center;
    gap: 7px;
    flex-wrap: wrap;
    margin-top: 5px;
}

.feature-badge {
    padding: 4px 9px;
    border: 1px solid #d9ddff;
    border-radius: 999px;
    background: #f7f7ff;
    color: #4944b8;
    font-size: 12px;
    font-weight: 750;
}

.workspace {
    background: var(--surface);
    border: 1px solid #edf0f7;
    border-radius: 24px;
    padding: 24px;
    box-shadow: 0 18px 45px rgba(30, 41, 59, 0.09);
}

.workspace-heading {
    margin-bottom: 18px;
}

.workspace-heading strong {
    color: var(--text);
    font-size: 18px;
}

.workspace-heading span {
    display: block;
    margin-top: 4px;
    color: var(--muted);
    font-size: 13px;
}

#top-row,
#bottom-row {
    gap: 22px !important;
    align-items: stretch !important;
}

#bottom-row {
    margin-top: 22px !important;
}

.section-card {
    border: 1px solid var(--border);
    border-radius: 18px;
    padding: 16px !important;
    background: #fbfcff;
    height: 100%;
}

.section-label {
    display: flex;
    align-items: center;
    gap: 8px;
    margin-bottom: 12px;
    color: var(--text);
    font-size: 15px;
    font-weight: 800;
}

.step-number {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 27px;
    height: 27px;
    border-radius: 9px;
    background: var(--primary);
    color: white;
    font-size: 13px;
    font-weight: 850;
}

.step-icon {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 26px;
    height: 26px;
    border-radius: 8px;
    background: #eef2ff;
    border: 1px solid #d9ddff;
    font-size: 14px;
    line-height: 1;
    flex: 0 0 auto;
}

.helper-text {
    color: var(--muted);
    font-size: 12px;
    line-height: 1.45;
    margin: -4px 0 10px 0;
}

#pdf-upload {
    min-height: 210px !important;
    max-height: 210px !important;
    border: 2px dashed #8b86f5 !important;
    border-radius: 15px !important;
    background: white !important;
    overflow-y: auto !important;
}

#pdf-upload > div {
    min-height: 205px !important;
}

.export-action-row {
    align-items: center !important;
    gap: 10px !important;
    margin-top: 10px !important;
}

.copy-status {
    min-height: 22px;
    color: #166534;
    font-size: 12px;
}

.copy-status.warning {
    color: #92400e;
}

.export-card {
    margin-top: 16px !important;
}

.export-helper {
    color: var(--muted);
    font-size: 13px;
    line-height: 1.5;
    margin-bottom: 10px;
}

.export-success {
    margin-top: 10px;
}

.action-row {
    gap: 10px !important;
    margin-top: 10px !important;
}

#process-button,
#ask-button,
#clear-button {
    min-height: 43px !important;
    border-radius: 12px !important;
    font-size: 14px !important;
    font-weight: 800 !important;
}

#process-button,
#ask-button {
    box-shadow: 0 7px 16px rgba(88, 80, 236, 0.24);
}

#clear-button {
    border: 1px solid #ccd3e4 !important;
    background: white !important;
    color: #475569 !important;
}

.status-box {
    margin-top: 12px;
    border-radius: 14px;
    padding: 13px;
    font-size: 13px;
    line-height: 1.45;
}

.status-title {
    font-size: 14px;
    font-weight: 850;
    margin-bottom: 8px;
}

.status-metrics {
    display: flex;
    flex-wrap: wrap;
    gap: 7px;
    margin-bottom: 9px;
}

.status-metrics span {
    background: rgba(255, 255, 255, 0.70);
    border-radius: 8px;
    padding: 5px 8px;
}

.success {
    background: #eafbf0;
    color: #17603a;
    border: 1px solid #72d39b;
}

.error {
    background: #fff0f0;
    color: #9f2727;
    border: 1px solid #f0a0a0;
}

.document-results {
    display: grid;
    gap: 5px;
}

.document-row {
    display: flex;
    justify-content: space-between;
    gap: 12px;
    border-top: 1px solid rgba(22, 101, 52, 0.12);
    padding-top: 6px;
}

.document-name {
    font-weight: 700;
    overflow-wrap: anywhere;
}

.page-count {
    white-space: nowrap;
    opacity: 0.82;
}

.warning-list {
    margin-top: 10px;
}

.summary-box,
.answer-box {
    height: 100%;
    min-height: 335px;
    max-height: 520px;
    overflow-y: auto;
    box-sizing: border-box;
    background: white;
    border: 1px solid #ccd3ff;
    border-radius: 15px;
    padding: 15px;
    color: var(--text);
    box-shadow: 0 5px 14px rgba(88, 80, 236, 0.07);
}

.panel-heading {
    position: sticky;
    top: -15px;
    z-index: 1;
    margin: -15px -15px 13px -15px;
    padding: 12px 15px;
    background: #f0f2ff;
    border-bottom: 1px solid #d6d9ff;
    color: #38379c;
    font-weight: 850;
}

.panel-content,
.answer-text {
    font-size: 14px;
    line-height: 1.58;
}

.formatted-ai-text p {
    margin: 0 0 10px 0;
}

.formatted-ai-text h3,
.formatted-ai-text h4,
.formatted-ai-text h5 {
    margin: 15px 0 8px 0;
    color: #2f3586;
    line-height: 1.3;
}

.formatted-ai-text h3 {
    font-size: 17px;
}

.formatted-ai-text h4 {
    font-size: 15px;
}

.formatted-ai-text h5 {
    font-size: 14px;
}

.formatted-ai-text ul,
.formatted-ai-text ol {
    margin: 7px 0 12px 21px;
    padding: 0;
}

.formatted-ai-text li {
    margin-bottom: 6px;
}

.formatted-ai-text strong {
    color: #202944;
}

.text-spacer {
    height: 4px;
}

.empty-panel {
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    text-align: center;
    gap: 7px;
    color: var(--muted);
}

.empty-panel strong {
    color: var(--text);
    font-size: 16px;
}

.empty-panel span {
    max-width: 390px;
    line-height: 1.45;
    font-size: 13px;
}

.empty-icon {
    font-size: 28px;
}

#question-box textarea {
    min-height: 150px !important;
    max-height: 150px !important;
    background: white !important;
    border: 1px solid #9d99f6 !important;
    border-radius: 14px !important;
    color: var(--text) !important;
    font-size: 14px !important;
    line-height: 1.5 !important;
    box-shadow: 0 0 0 3px rgba(88, 80, 236, 0.06);
}

.question-detection {
    margin-bottom: 12px;
    padding: 9px 11px;
    border-radius: 10px;
    background: #ecfdf3;
    border: 1px solid #a5dfbd;
    color: #166534;
    font-size: 13px;
    font-weight: 750;
}

.search-interpretation {
    margin: 0 0 10px 0;
    padding: 8px 10px;
    border-radius: 9px;
    background: #eef6ff;
    border: 1px solid #bdd7f6;
    color: #245985;
    font-size: 12px;
    font-weight: 700;
}

.individual-answer {
    padding: 3px 0 15px 0;
}

.individual-answer + .individual-answer {
    border-top: 1px solid #e1e5f0;
    padding-top: 16px;
}

.question-heading {
    display: flex;
    align-items: flex-start;
    gap: 8px;
    margin-bottom: 9px;
    color: #30358d;
    font-size: 14px;
    font-weight: 850;
}

.answer-number {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    min-width: 24px;
    height: 24px;
    border-radius: 8px;
    background: var(--primary);
    color: white;
    font-size: 12px;
}

.sources-box {
    margin-top: 13px;
    padding: 11px;
    border: 1px solid #dbe0f0;
    border-radius: 11px;
    background: #f8f9fd;
}

.sources-title {
    color: #46506a;
    font-size: 12px;
    font-weight: 850;
    margin-bottom: 7px;
}

.sources-box ul {
    list-style: none;
    padding: 0;
    margin: 0;
    display: grid;
    gap: 6px;
}

.sources-box li {
    display: flex;
    justify-content: space-between;
    gap: 12px;
    font-size: 12px;
}

.source-file {
    color: #343b52;
    overflow-wrap: anywhere;
}

.source-page {
    color: #626b82;
    white-space: nowrap;
}

.service-warning-panel {
    display: block !important;
}

.service-warning {
    display: flex;
    flex-direction: column;
    gap: 7px;
    margin: 8px 0;
    padding: 14px;
    border: 1px solid #f2c66d;
    border-radius: 11px;
    background: #fff8e8;
    color: #7a4d0b;
    font-size: 13px;
    line-height: 1.5;
}

.service-warning strong {
    font-size: 14px;
}

.quick-review-success {
    margin-bottom: 13px;
    padding: 9px 11px;
    border: 1px solid #a5dfbd;
    border-radius: 10px;
    background: #ecfdf3;
    color: #166534;
    font-size: 13px;
    font-weight: 750;
}

.quick-review-heading {
    display: flex;
    align-items: baseline;
    gap: 8px;
    margin: 12px 0 8px 0;
    color: #30358d;
    font-size: 13px;
    font-weight: 850;
}

.quick-review-heading span {
    color: var(--muted);
    font-size: 11px;
    font-weight: 500;
}

.quick-review-row {
    gap: 8px !important;
    margin-bottom: 7px !important;
}

.quick-review-button {
    min-height: 38px !important;
    border: 1px solid #cfd4ff !important;
    border-radius: 10px !important;
    background: #f7f7ff !important;
    color: #3f3c9e !important;
    font-size: 12px !important;
    font-weight: 750 !important;
}

.quick-review-button:hover {
    background: #ececff !important;
    border-color: #9490f3 !important;
}

.safety-note {
    margin-top: 17px;
    padding: 11px 13px;
    border-radius: 12px;
    background: #fffaf0;
    border: 1px solid #f2d49b;
    color: #795117;
    font-size: 12px;
    line-height: 1.45;
}


.structured-review-summary {
    display: flex;
    flex-direction: column;
    gap: 5px;
    margin: 12px;
    padding: 11px 13px;
    border-radius: 11px;
    background: #f7f8ff;
    border: 1px solid #dcdcff;
    color: #34367d;
    font-size: 13px;
    line-height: 1.45;
}

.structured-table-wrap {
    width: calc(100% - 24px);
    margin: 12px;
    overflow-x: auto;
    border: 1px solid #dbe2f2;
    border-radius: 12px;
    background: #ffffff;
}

.structured-review-table {
    width: 100%;
    min-width: 760px;
    border-collapse: collapse;
    font-size: 12px;
    line-height: 1.45;
}

.structured-review-table th {
    padding: 10px;
    text-align: left;
    vertical-align: top;
    background: #eef2ff;
    color: #37358d;
    border-bottom: 1px solid #dbe2f2;
    font-weight: 800;
    white-space: nowrap;
}

.structured-review-table td {
    padding: 10px;
    vertical-align: top;
    color: #273248;
    border-bottom: 1px solid #edf0f7;
    word-break: break-word;
}

.structured-review-table tbody tr:last-child td {
    border-bottom: none;
}

.structured-review-table tbody tr:nth-child(even) {
    background: #fafbff;
}

.severity-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    min-width: 58px;
    padding: 3px 8px;
    border-radius: 999px;
    font-size: 11px;
    font-weight: 800;
    border: 1px solid transparent;
}

.severity-high {
    background: #fff0f0;
    color: #b42318;
    border-color: #ffc9c9;
}

.severity-medium {
    background: #fff8e6;
    color: #9a5b00;
    border-color: #f7d78b;
}

.severity-low {
    background: #edfdf3;
    color: #18794e;
    border-color: #afe1c3;
}

.structured-empty,
.structured-fallback-note {
    margin: 12px;
    padding: 12px 13px;
    border-radius: 11px;
    background: #fffaf0;
    border: 1px solid #f2d49b;
    color: #795117;
    font-size: 12px;
    line-height: 1.45;
}

@media (max-width: 800px) {
    .gradio-container {
        width: 100vw !important;
        height: 100dvh !important;
        max-height: 100dvh !important;
        padding-top: 0 !important;
    }

    .main-title {
        font-size: 27px;
    }

    .workspace {
        padding: 14px;
    }

    .summary-box,
    .answer-box {
        min-height: 270px;
    }

    .document-row,
    .sources-box li {
        flex-direction: column;
        gap: 2px;
    }

    .page-count,
    .source-page {
        white-space: normal;
    }


    #sticky-header {
        width: 100% !important;
    }
}
"""


# ---------------------------------------------------------
# Header / scroll JavaScript fallbacks
# ---------------------------------------------------------

header_scroll_js = r"""
() => {
    const boot = () => {
        const shell = document.getElementById("finops-scroll-shell");
        const header = document.getElementById("finops-topbar");

        if (!shell || !header) {
            window.setTimeout(boot, 120);
            return;
        }

        const measure = () => {
            const shellRect = shell.getBoundingClientRect();
            const headerRect = header.getBoundingClientRect();

            shell.style.setProperty("--finops-shell-left", `${Math.round(shellRect.left)}px`);
            shell.style.setProperty("--finops-shell-width", `${Math.round(shellRect.width)}px`);
            shell.style.setProperty("--finops-topbar-height", `${Math.ceil(headerRect.height)}px`);
        };

        const verifySticky = () => {
            measure();

            // Only test once the user has actually scrolled inside the FinOps workspace.
            if (shell.scrollTop < 24) {
                return;
            }

            const shellTop = shell.getBoundingClientRect().top;
            const headerTop = header.getBoundingClientRect().top;

            // Sticky should keep the two tops essentially aligned.
            // If not, automatically switch to the fixed fallback.
            if (Math.abs(headerTop - shellTop) > 4) {
                shell.classList.add("finops-fixed-fallback");
                measure();
            }
        };

        // Ensure the app starts at the top of its own scroll area.
        shell.scrollTop = 0;

        shell.addEventListener("scroll", verifySticky, { passive: true });
        window.addEventListener("resize", measure, { passive: true });

        if ("ResizeObserver" in window) {
            const ro = new ResizeObserver(measure);
            ro.observe(shell);
            ro.observe(header);
        }

        measure();
        window.setTimeout(measure, 300);
        window.setTimeout(measure, 1000);
    };

    boot();
}
"""


def processing_html(message):
    """Return a clear processing banner shown while a long-running action executes."""
    safe_message = html.escape(message)
    return f"""
    <div class="finops-processing" role="status" aria-live="polite">
        <div class="finops-processing-spinner" aria-hidden="true"></div>
        <div class="finops-processing-copy">
            <strong>{safe_message}</strong>
            <span>Please wait — the Copilot is working on your documents.</span>
        </div>
    </div>
    """


# ---------------------------------------------------------
# Build Gradio interface
# ---------------------------------------------------------

with gr.Blocks(
    title="FinOps AI Copilot v0.10",
    fill_width=True,
) as demo:

    chunks_state = gr.State([])
    embeddings_state = gr.State(None)
    review_results_state = gr.State({})
    summary_text_state = gr.State("")
    processed_documents_state = gr.State([])
    qa_history_state = gr.State([])
    current_copy_text_state = gr.State("")
    current_review_key_state = gr.State("")
    ai_recommendation_state = gr.State("")
    hitl_decision_state = gr.State("")


    # v0.8: Header and all visible content share ONE scroll container.
    with gr.Column(elem_id="finops-scroll-shell"):

        gr.HTML(
            """
            <div class="app-hero">
                <div class="title-line">
                    <span class="main-title">💼 FinOps AI Copilot</span>
                    <span class="version-badge">MVP v0.10</span>
                </div>
                <div class="sub-title">
                    Review expense claims against finance policy, surface exceptions and missing evidence,
                    then keep the final decision with a human Finance reviewer.
                </div>
                <div class="feature-strip">
                    <span class="feature-badge">🧾 Expense review</span>
                    <span class="feature-badge">📑 Policy compliance</span>
                    <span class="feature-badge">⚠️ Exceptions & evidence</span>
                    <span class="feature-badge">👤 Human-in-the-Loop</span>
                    <span class="feature-badge">📄 Finance report</span>
                </div>
            </div>
            """,
            elem_id="finops-topbar",
        )

        gr.HTML('<div id="finops-topbar-spacer" aria-hidden="true"></div>')

        with gr.Column(elem_classes="workspace"):
            gr.HTML(
                """
                <div class="workspace-heading">
                    <strong>FinOps review workspace</strong>
                    <span>
                        Upload expense and policy documents, run the finance checks, review the AI recommendation,
                        then record the final human decision.
                    </span>
                </div>
                """
            )

            with gr.Row(elem_id="top-row"):
                with gr.Column(scale=1, elem_classes="section-card"):
                    gr.HTML(
                        """
                        <div class="section-label">
                            <span class="step-icon" title="Upload documents">📤</span>
                            <span class="step-number">1</span>
                            <span>Upload expense & policy documents</span>
                        </div>
                        <div class="helper-text">
                            Upload travel expense claims, receipts/invoices and finance or travel policy documents.
                            PDF, PNG, JPG/JPEG and WEBP are supported. Scanned/image documents automatically use Gemini vision/OCR fallback.
                        </div>
                        """
                    )

                    pdf_input = gr.File(
                        label=None,
                        show_label=False,
                        file_types=[".pdf", ".png", ".jpg", ".jpeg", ".webp"],
                        file_count="multiple",
                        type="filepath",
                        elem_id="pdf-upload",
                    )

                    with gr.Row(elem_classes="action-row"):
                        process_button = gr.Button(
                            "Process Finance Documents 🚀",
                            variant="primary",
                            elem_id="process-button",
                            scale=3,
                        )

                        clear_button = gr.Button(
                            "Clear All",
                            variant="secondary",
                            elem_id="clear-button",
                            scale=1,
                        )

                    status_output = gr.HTML()

                with gr.Column(scale=1, elem_classes="section-card"):
                    gr.HTML(
                        """
                        <div class="section-label">
                            <span class="step-icon" title="Review summary">🧾</span>
                            <span class="step-number">2</span>
                            <span>Review finance case summary</span>
                        </div>
                        <div class="helper-text">
                            Gemini creates a grounded finance summary from the readable text
                            across the uploaded claim, evidence and policy documents.
                        </div>
                        """
                    )

                    summary_output = gr.HTML(EMPTY_SUMMARY_HTML, elem_id="summary-copy-source")

                    with gr.Row(elem_classes="export-action-row"):
                        copy_summary_button = gr.Button(
                            "📋 Copy Summary",
                            variant="secondary",
                            size="sm",
                        )
                        summary_copy_status = gr.HTML()

            with gr.Row(elem_id="bottom-row"):
                with gr.Column(scale=1, elem_classes="section-card"):
                    gr.HTML(
                        """
                        <div class="section-label">
                            <span class="step-icon" title="Ask questions">💬</span>
                            <span class="step-number">3</span>
                            <span>Ask finance / policy questions</span>
                        </div>
                        <div class="helper-text">
                            Ask about an expense, policy rule, exception or missing evidence.
                            You can also ask several questions in one message.
                        </div>
                        """
                    )

                    question_input = gr.Textbox(
                        label=None,
                        show_label=False,
                        placeholder=(
                            "Example: Is the hotel claim within policy, and what evidence is missing?"
                        ),
                        lines=5,
                        elem_id="question-box",
                    )

                    gr.HTML(
                        """
                        <div class="quick-review-heading">
                            ⚡ Finance Quick Review
                            <span>Run the four core finance checks with one click.</span>
                        </div>
                        """
                    )

                    with gr.Row(elem_classes="quick-review-row"):
                        risks_button = gr.Button(
                            "🧾 Expense Summary",
                            variant="secondary",
                            elem_classes="quick-review-button",
                        )

                        actions_button = gr.Button(
                            "📑 Policy Compliance",
                            variant="secondary",
                            elem_classes="quick-review-button",
                        )

                    with gr.Row(elem_classes="quick-review-row"):
                        dates_button = gr.Button(
                            "⚠️ Exceptions & Risks",
                            variant="secondary",
                            elem_classes="quick-review-button",
                        )

                        compare_button = gr.Button(
                            "📎 Missing Evidence",
                            variant="secondary",
                            elem_classes="quick-review-button",
                        )

                    ask_button = gr.Button(
                        "Ask Finance Documents 💬",
                        variant="primary",
                        elem_id="ask-button",
                    )

                    gr.HTML(
                        """
                        <div class="safety-note">
                            <strong>Important:</strong> AI can make mistakes. Verify material findings against the cited pages.
                            The AI recommendation is advisory; Finance makes the final decision.
                        </div>
                        """
                    )

                with gr.Column(scale=1, elem_classes="section-card"):
                    gr.HTML(
                        """
                        <div class="section-label">
                            <span class="step-icon" title="Check findings">🔎</span>
                            <span class="step-number">4</span>
                            <span>Check findings and evidence</span>
                        </div>
                        <div class="helper-text">
                            Structured findings show the supporting file and page where available.
                            Use these references before making the final finance decision.
                        </div>
                        """
                    )

                    answer_output = gr.HTML(EMPTY_ANSWER_HTML, elem_id="result-copy-source")

                    with gr.Row(elem_classes="export-action-row"):
                        copy_result_button = gr.Button(
                            "📋 Copy Current Result",
                            variant="secondary",
                            size="sm",
                        )
                        csv_download_button = gr.DownloadButton(
                            "⬇️ Download Current Finance Table CSV",
                            value=generate_current_review_csv,
                            inputs=[
                                current_review_key_state,
                                review_results_state,
                            ],
                            variant="secondary",
                            size="sm",
                        )

                    result_copy_status = gr.HTML()


            with gr.Column(elem_classes=["section-card", "export-card"]):
                gr.HTML(
                    """
                    <div class="section-label">
                        <span class="step-icon" title="AI and human review">🤝</span>
                        <span class="step-number">5</span>
                        <span>AI recommendation & Human-in-the-Loop review</span>
                    </div>
                    <div class="export-helper">
                        Generate an advisory AI recommendation, then Finance records the final decision.
                        The AI does not approve expenses by itself.
                    </div>
                    """
                )

                generate_recommendation_button = gr.Button(
                    "🤖 Generate AI Recommendation",
                    variant="primary",
                )
                ai_recommendation_output = gr.HTML(EMPTY_RECOMMENDATION_HTML)

                reviewer_note = gr.Textbox(
                    label="Reviewer note",
                    placeholder="Optional: reason for approval/rejection/escalation or evidence to request",
                    lines=3,
                )

                with gr.Row(elem_classes="quick-review-row"):
                    approve_button = gr.Button("✅ Approve", variant="secondary")
                    reject_button = gr.Button("❌ Reject", variant="secondary")
                    escalate_button = gr.Button("🧑‍💼 Escalate / Manual Review", variant="secondary")

                hitl_status_output = gr.HTML()

            with gr.Column(elem_classes=["section-card", "export-card"]):
                gr.HTML(
                    """
                    <div class="section-label">
                        <span class="step-icon" title="Export report">📄</span>
                        <span class="step-number">6</span>
                        <span>Export the Finance Review Report</span>
                    </div>
                    <div class="export-helper">
                        Generate an editable Word report containing the finance case summary,
                        all four structured finance checks, source/page evidence, AI recommendation,
                        the human Finance decision and reviewer notes. Missing review sections are
                        generated automatically before export.
                    </div>
                    """
                )

                generate_report_button = gr.Button(
                    "📄 Generate Finance Review Report",
                    variant="primary",
                    elem_id="report-button",
                )
                report_status_output = gr.HTML()
                report_file_output = gr.File(
                    label="Download Finance Review Report",
                    interactive=False,
                )

    # ---------------------------------------------------------
    # Long-running actions with clear processing feedback
    # ---------------------------------------------------------

    process_start = process_button.click(
        fn=lambda: processing_html("Processing finance documents…"),
        inputs=[],
        outputs=summary_output,
        queue=False,
        show_progress="hidden",
    )

    process_event = process_start.then(
        fn=process_documents,
        inputs=pdf_input,
        show_progress="hidden",
        outputs=[
            status_output,
            summary_output,
            chunks_state,
            embeddings_state,
            summary_text_state,
            processed_documents_state,
        ],
    )

    process_event.then(
        fn=lambda: ({}, [], "", "", None, "", "", "", EMPTY_RECOMMENDATION_HTML, "", "", "", ""),
        inputs=[],
        show_progress="hidden",
        outputs=[
            review_results_state,
            qa_history_state,
            current_copy_text_state,
            current_review_key_state,
            report_file_output,
            report_status_output,
            summary_copy_status,
            result_copy_status,
            ai_recommendation_output,
            ai_recommendation_state,
            hitl_decision_state,
            reviewer_note,
            hitl_status_output,
        ],
    )

    ask_start = ask_button.click(
        fn=lambda: processing_html("Searching finance documents and preparing an answer…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    ask_start.then(
        fn=answer_question,
        show_progress="hidden",
        inputs=[
            question_input,
            chunks_state,
            embeddings_state,
            qa_history_state,
        ],
        outputs=[
            answer_output,
            qa_history_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    submit_start = question_input.submit(
        fn=lambda: processing_html("Searching finance documents and preparing an answer…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    submit_start.then(
        fn=answer_question,
        show_progress="hidden",
        inputs=[
            question_input,
            chunks_state,
            embeddings_state,
            qa_history_state,
        ],
        outputs=[
            answer_output,
            qa_history_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    risks_start = risks_button.click(
        fn=lambda: processing_html("Reviewing exceptions and risks…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    risks_start.then(
        fn=review_key_risks,
        show_progress="hidden",
        inputs=[
            chunks_state,
            embeddings_state,
            review_results_state,
        ],
        outputs=[
            question_input,
            answer_output,
            review_results_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    actions_start = actions_button.click(
        fn=lambda: processing_html("Checking policy compliance…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    actions_start.then(
        fn=review_actions_obligations,
        show_progress="hidden",
        inputs=[
            chunks_state,
            embeddings_state,
            review_results_state,
        ],
        outputs=[
            question_input,
            answer_output,
            review_results_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    dates_start = dates_button.click(
        fn=lambda: processing_html("Extracting the expense summary…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    dates_start.then(
        fn=review_important_dates,
        show_progress="hidden",
        inputs=[
            chunks_state,
            embeddings_state,
            review_results_state,
        ],
        outputs=[
            question_input,
            answer_output,
            review_results_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    compare_start = compare_button.click(
        fn=lambda: processing_html("Checking for missing evidence…"),
        inputs=[],
        outputs=answer_output,
        queue=False,
        show_progress="hidden",
    )

    compare_start.then(
        fn=review_document_comparison,
        show_progress="hidden",
        inputs=[
            chunks_state,
            embeddings_state,
            review_results_state,
        ],
        outputs=[
            question_input,
            answer_output,
            review_results_state,
            current_copy_text_state,
            current_review_key_state,
        ],
    )

    recommendation_start = generate_recommendation_button.click(
        fn=lambda: processing_html("Generating the advisory AI recommendation…"),
        inputs=[],
        outputs=ai_recommendation_output,
        queue=False,
        show_progress="hidden",
    )

    recommendation_start.then(
        fn=generate_finance_recommendation,
        show_progress="hidden",
        inputs=[chunks_state, review_results_state],
        outputs=[
            ai_recommendation_output,
            ai_recommendation_state,
            review_results_state,
        ],
    )

    # HITL is intentionally unchanged: the human decision is immediate and explicit.
    approve_button.click(
        fn=hitl_approve,
        show_progress="hidden",
        inputs=[reviewer_note, ai_recommendation_state],
        outputs=[hitl_status_output, hitl_decision_state],
    )

    reject_button.click(
        fn=hitl_reject,
        show_progress="hidden",
        inputs=[reviewer_note, ai_recommendation_state],
        outputs=[hitl_status_output, hitl_decision_state],
    )

    escalate_button.click(
        fn=hitl_escalate,
        show_progress="hidden",
        inputs=[reviewer_note, ai_recommendation_state],
        outputs=[hitl_status_output, hitl_decision_state],
    )


    # Copy directly from the visible HTML output instead of reading gr.State.
    # This avoids stale/empty State values in browser-only JavaScript events.
    copy_summary_button.click(
        fn=None,
        inputs=[],
        outputs=[],
        queue=False,
        show_progress="hidden",
        js="""async () => {
            const source = document.getElementById('summary-copy-source');

            if (!source) {
                window.alert('The summary area could not be found. Please refresh the page and try again.');
                return [];
            }

            const clone = source.cloneNode(true);
            clone.querySelectorAll('script, style, button').forEach((node) => node.remove());

            clone.querySelectorAll('table').forEach((table) => {
                const lines = Array.from(table.rows).map((row) =>
                    Array.from(row.cells)
                        .map((cell) => cell.innerText.trim().replace(/\\s+/g, ' '))
                        .join('\\t')
                );
                const replacement = document.createElement('pre');
                replacement.textContent = lines.join('\\n');
                table.replaceWith(replacement);
            });

            const value = clone.innerText
                .replace(/\\n{3,}/g, '\\n\\n')
                .trim();

            if (!value || value.includes('No summary yet')) {
                window.alert('Nothing to copy yet. Process finance documents first.');
                return [];
            }

            const fallbackCopy = () => {
                const area = document.createElement('textarea');
                area.value = value;
                area.style.position = 'fixed';
                area.style.left = '-9999px';
                area.style.top = '0';
                document.body.appendChild(area);
                area.focus();
                area.select();
                area.setSelectionRange(0, area.value.length);
                const copied = document.execCommand('copy');
                document.body.removeChild(area);
                return copied;
            };

            try {
                if (navigator.clipboard && navigator.clipboard.writeText) {
                    await navigator.clipboard.writeText(value);
                } else if (!fallbackCopy()) {
                    throw new Error('Clipboard copy failed');
                }
            } catch (error) {
                if (!fallbackCopy()) {
                    window.alert('The browser blocked clipboard access. Please allow clipboard permission and try again.');
                    return [];
                }
            }

            const toast = document.createElement('div');
            toast.textContent = 'Summary copied to clipboard';
            toast.style.position = 'fixed';
            toast.style.right = '24px';
            toast.style.bottom = '24px';
            toast.style.zIndex = '99999';
            toast.style.padding = '12px 16px';
            toast.style.borderRadius = '10px';
            toast.style.background = '#111827';
            toast.style.color = '#ffffff';
            toast.style.fontWeight = '600';
            toast.style.boxShadow = '0 8px 24px rgba(0,0,0,0.22)';
            document.body.appendChild(toast);
            window.setTimeout(() => toast.remove(), 1800);
            return [];
        }""",
    )

    copy_result_button.click(
        fn=None,
        inputs=[],
        outputs=[],
        queue=False,
        show_progress="hidden",
        js="""async () => {
            const source = document.getElementById('result-copy-source');

            if (!source) {
                window.alert('The result area could not be found. Please refresh the page and try again.');
                return [];
            }

            const clone = source.cloneNode(true);
            clone.querySelectorAll('script, style, button').forEach((node) => node.remove());

            clone.querySelectorAll('table').forEach((table) => {
                const lines = Array.from(table.rows).map((row) =>
                    Array.from(row.cells)
                        .map((cell) => cell.innerText.trim().replace(/\\s+/g, ' '))
                        .join('\\t')
                );
                const replacement = document.createElement('pre');
                replacement.textContent = lines.join('\\n');
                table.replaceWith(replacement);
            });

            const value = clone.innerText
                .replace(/\\n{3,}/g, '\\n\\n')
                .trim();

            if (!value || value.includes('No answer yet')) {
                window.alert('Run a review or ask a question first.');
                return [];
            }

            const fallbackCopy = () => {
                const area = document.createElement('textarea');
                area.value = value;
                area.style.position = 'fixed';
                area.style.left = '-9999px';
                area.style.top = '0';
                document.body.appendChild(area);
                area.focus();
                area.select();
                area.setSelectionRange(0, area.value.length);
                const copied = document.execCommand('copy');
                document.body.removeChild(area);
                return copied;
            };

            try {
                if (navigator.clipboard && navigator.clipboard.writeText) {
                    await navigator.clipboard.writeText(value);
                } else if (!fallbackCopy()) {
                    throw new Error('Clipboard copy failed');
                }
            } catch (error) {
                if (!fallbackCopy()) {
                    window.alert('The browser blocked clipboard access. Please allow clipboard permission and try again.');
                    return [];
                }
            }

            const toast = document.createElement('div');
            toast.textContent = 'Current result copied to clipboard';
            toast.style.position = 'fixed';
            toast.style.right = '24px';
            toast.style.bottom = '24px';
            toast.style.zIndex = '99999';
            toast.style.padding = '12px 16px';
            toast.style.borderRadius = '10px';
            toast.style.background = '#111827';
            toast.style.color = '#ffffff';
            toast.style.fontWeight = '600';
            toast.style.boxShadow = '0 8px 24px rgba(0,0,0,0.22)';
            document.body.appendChild(toast);
            window.setTimeout(() => toast.remove(), 1800);
            return [];
        }""",
    )

    report_start = generate_report_button.click(
        fn=lambda: processing_html("Generating the Finance Review Report…"),
        inputs=[],
        outputs=report_status_output,
        queue=False,
        show_progress="hidden",
    )

    report_start.then(
        fn=generate_complete_report,
        show_progress="hidden",
        inputs=[
            chunks_state,
            embeddings_state,
            summary_text_state,
            processed_documents_state,
            review_results_state,
            qa_history_state,
            ai_recommendation_state,
            hitl_decision_state,
            reviewer_note,
        ],
        outputs=[
            report_file_output,
            report_status_output,
            review_results_state,
        ],
    )

    clear_button.click(
        fn=clear_all,
        show_progress="hidden",
        inputs=[],
        outputs=[
            pdf_input,
            status_output,
            summary_output,
            question_input,
            answer_output,
            chunks_state,
            embeddings_state,
            review_results_state,
            summary_text_state,
            processed_documents_state,
            qa_history_state,
            current_copy_text_state,
            current_review_key_state,
            summary_copy_status,
            result_copy_status,
            report_file_output,
            report_status_output,
            ai_recommendation_output,
            ai_recommendation_state,
            hitl_decision_state,
            reviewer_note,
            hitl_status_output,
        ],
    )


demo.launch(
    server_name="0.0.0.0",
    server_port=7860,
    theme=gr.themes.Soft(),
    css=css,
    js=header_scroll_js,
)
